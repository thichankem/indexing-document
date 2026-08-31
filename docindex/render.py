"""Write the computed layout out as a .docx or .pdf file.

Both writers read the same list of `LayoutPage` objects, so the two formats
produce the same layout: headings alone on their line, sizes shrinking with
depth.

The presentation here exists to serve the DLA model that will read the document
back. That model looks at the printed page rather than the file structure, so
the label it assigns each line depends on exactly three geometric cues:

* **title** — bold, clearly larger than the body, alone on its line against the
  left margin, with generous whitespace above and below;
* **list-item** — a line opening with a bullet marker, indented from the
  margin, at body font size;
* **text** — a justified paragraph, no indent, no marker.

Hence headings get wide spacing and keep their size-per-level, and bullet lines
get a real indent — look like a properly formatted administrative document and
the DLA model builds the right hierarchy.
"""
from __future__ import annotations

import html
import io
import os
import re

import docx
import fitz
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .layout import BODY_PT, LayoutPage, PageItem

# Markers that open a list line in Vietnamese administrative documents
_BULLET = re.compile(r"^\s*([-–—•+*]|\(?[a-zđ]\)|\(?[ivx]+\)|\d+[.)])\s+")

# Whitespace around a heading, in points. Wider than the gap between paragraphs
# so the DLA model sees the heading as its own block, not the first line of a
# paragraph.
HEAD_SPACE_BEFORE = 16.0
HEAD_SPACE_AFTER = 9.0
LIST_INDENT_PT = 18.0          # indent of a bullet line

# A gap narrower than this cannot hold even one line of text
MIN_PLACE_PT = 24.0
# Guard on the block-placement loop: no block is ever this many pages long
_MAX_PAGES_PER_BLOCK = 500
# Rebuild passes used to push headings stranded at the foot of a page onto the
# next one. Moving a heading down shifts everything below it and sometimes
# exposes another stranded heading, so several passes may be needed; the loop
# stops as soon as a pass finds nothing new, so an ordinary document costs one.
_MAX_REFLOW_PASSES = 10

# Width of the A4 content area at 2cm margins, in points; used to fit images
CONTENT_WIDTH_PT = 470.0

# The font must carry the full set of Vietnamese diacritics. The PDF base-14
# fonts do not, so a system font has to be embedded; the first pair found wins.
_FONT_CANDIDATES = [
    ("times.ttf", "timesbd.ttf"),
    ("arial.ttf", "arialbd.ttf"),
    ("segoeui.ttf", "segoeuib.ttf"),
    ("DejaVuSans.ttf", "DejaVuSans-Bold.ttf"),
]
_FONT_DIRS = ["C:/Windows/Fonts", "/usr/share/fonts/truetype/dejavu",
              "/Library/Fonts", "/usr/share/fonts"]

DOCX_FONT = "Times New Roman"


# --- .docx ---------------------------------------------------------------

def _md_rows(md: str) -> list[list[str]]:
    """Parse a markdown table back into rows of cells."""
    rows: list[list[str]] = []
    for line in md.split("\n"):
        line = line.strip()
        if not line.startswith("|"):
            continue
        if set(line.replace("|", "").replace(" ", "")) <= {"-"}:
            continue                      # markdown's separator line
        cells = re.split(r"(?<!\\)\|", line.strip().strip("|"))
        rows.append([c.strip().replace("\\|", "|") for c in cells])
    return rows


def _setup_docx(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.space_after = Pt(4)
    for section in document.sections:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin = section.bottom_margin = Cm(2)


def _outline_level(para, level: int) -> None:
    """Mark this paragraph as a level-`level` heading in Word's own structure.

    Size and weight are what the DLA model sees; `outlineLvl` is what Word and
    the PDF converters read, which is how the exported PDF ends up with a
    bookmark tree at the right levels.
    """
    pPr = para._p.get_or_add_pPr()
    tag = OxmlElement("w:outlineLvl")
    tag.set(qn("w:val"), str(min(max(level, 1), 9) - 1))
    pPr.append(tag)


def _docx_heading(document, item: PageItem) -> None:
    para = document.add_paragraph()
    # Generous whitespace above and below is the strongest cue the DLA model has
    # for separating a heading from the paragraph next to it.
    para.paragraph_format.space_before = Pt(HEAD_SPACE_BEFORE)
    para.paragraph_format.space_after = Pt(HEAD_SPACE_AFTER)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.keep_together = True
    # The document title is centred as in any administrative document; section
    # headings are left-aligned
    para.alignment = (WD_ALIGN_PARAGRAPH.CENTER if item.level <= 0
                      else WD_ALIGN_PARAGRAPH.LEFT)
    _outline_level(para, max(item.level, 1))
    run = para.add_run(item.text)
    run.bold = True
    run.font.size = Pt(item.size)
    run.font.name = DOCX_FONT


def _docx_body(document, item: PageItem) -> None:
    """Write one content block, keeping bullet lines separate.

    Collapsing the block into a single justified paragraph makes bullet lines
    look exactly like prose; a real indent gets them labelled `list-item` and the
    hierarchy keeps the parent-section / child-item relationship.
    """
    for line in item.text.split("\n"):
        if not line.strip():
            continue
        bullet = bool(_BULLET.match(line))
        para = document.add_paragraph()
        fmt = para.paragraph_format
        if bullet:
            fmt.left_indent = Pt(LIST_INDENT_PT)
            fmt.first_line_indent = Pt(-LIST_INDENT_PT / 2)
            fmt.space_after = Pt(2)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(line.strip())
        run.font.size = Pt(item.size)
        run.font.name = DOCX_FONT


def _row_property(row, tag: str) -> None:
    """Turn on a boolean row property (w:tblHeader, w:cantSplit)."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement(tag))


def _repeat_header_row(table) -> None:
    if table.rows:
        _row_property(table.rows[0], "w:tblHeader")


def _keep_row_together(row) -> None:
    _row_property(row, "w:cantSplit")


def _docx_table(document, item: PageItem) -> None:
    rows = _md_rows(item.text)
    if not rows:
        # no table rows left (it was spread into prose) -> write it as a
        # paragraph; never leave it blank, that would lose content
        _docx_body(document, item)
        return
    width = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    # A long table crossing a page break: repeat the header row on every page and
    # let no row be split in half — a split row is a fault someone has to correct
    # by hand downstream.
    _repeat_header_row(table)
    for row in table.rows:
        _keep_row_together(row)
    for r, cells in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            cell.text = cells[c] if c < len(cells) else ""
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(BODY_PT - 0.5)
                    run.font.name = DOCX_FONT
                    run.bold = r == 0
    document.add_paragraph()


def _docx_figure(document, item: PageItem) -> None:
    path = item.meta.get("file") or ""
    if path and os.path.isfile(path):
        width = float(item.meta.get("width") or 0) or CONTENT_WIDTH_PT
        document.add_picture(path, width=Pt(min(width, CONTENT_WIDTH_PT)))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = item.meta.get("caption") or ""
        if caption:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(caption)
            run.italic = True
            run.font.size = Pt(BODY_PT - 1)
        return
    # no image file -> keep the placeholder line so the figure leaves a trace
    para = document.add_paragraph()
    run = para.add_run(item.text)
    run.italic = True


def write_docx(pages: list[LayoutPage], dst: str) -> None:
    document = docx.Document()
    _setup_docx(document)

    for index, page in enumerate(pages):
        if index:
            document.add_page_break()
        for item in page.items:
            if item.kind == "heading":
                _docx_heading(document, item)
            elif item.kind == "table":
                _docx_table(document, item)
            elif item.kind == "figure":
                _docx_figure(document, item)
            else:
                _docx_body(document, item)

    os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)
    document.save(dst)


# --- .pdf ----------------------------------------------------------------

def _font_files() -> tuple[str, str, str] | None:
    """Find an available regular/bold font pair on this system."""
    for folder in _FONT_DIRS:
        if not os.path.isdir(folder):
            continue
        for regular, bold in _FONT_CANDIDATES:
            if os.path.isfile(os.path.join(folder, regular)):
                if not os.path.isfile(os.path.join(folder, bold)):
                    bold = regular
                return folder, regular, bold
    return None


def _css(font: tuple[str, str, str] | None) -> str:
    face = ""
    if font:
        _folder, regular, bold = font
        face = (f"@font-face {{font-family: doc; src: url({regular});}}\n"
                f"@font-face {{font-family: doc; src: url({bold}); font-weight: bold;}}\n"
                "* {font-family: doc;}\n")
    return face + f"""
body {{font-size: {BODY_PT}px;}}
p {{font-size: {BODY_PT}px; margin: 0 0 4px 0; text-align: justify;}}
p.li {{margin-left: {LIST_INDENT_PT}px; text-align: left;}}
p.fig {{text-align: center; font-style: italic;}}
h1, h2, h3, h4, h5 {{font-weight: bold; text-align: left; margin-left: 0;
        margin-top: {HEAD_SPACE_BEFORE}px; margin-bottom: {HEAD_SPACE_AFTER}px;}}
table {{border-collapse: collapse; width: 100%; margin-bottom: 6px;}}
td, th {{border: 1px solid #808080; padding: 2px 4px; font-size: {BODY_PT - 1}px;
        text-align: left; vertical-align: top;}}
th {{font-weight: bold;}}
"""


def _table_html(md: str) -> str:
    rows = _md_rows(md)
    if not rows:
        return ""
    out = ["<table>"]
    for index, cells in enumerate(rows):
        tag = "th" if index == 0 else "td"
        out.append("<tr>" + "".join(
            f"<{tag}>{html.escape(c)}</{tag}>" for c in cells) + "</tr>")
    out.append("</table>")
    return "".join(out)


def _item_html(item: PageItem) -> str:
    if item.kind == "heading":
        level = min(max(item.level, 1), 5)
        style = f"font-size: {item.size}px"
        if item.level <= 0:             # the document title
            style += "; text-align: center"
        return f'<h{level} style="{style}">{html.escape(item.text)}</h{level}>'

    if item.kind == "table":
        html_table = _table_html(item.text)
        if html_table:
            return html_table
        return f"<p>{html.escape(item.text)}</p>"   # already spread into prose

    if item.kind == "figure":
        path = item.meta.get("file") or ""
        if path and os.path.isfile(path):
            width = min(float(item.meta.get("width") or 0) or CONTENT_WIDTH_PT,
                        CONTENT_WIDTH_PT)
            out = (f'<p class="fig"><img src="{os.path.basename(path)}" '
                   f'width="{round(width)}"></p>')
            caption = item.meta.get("caption") or ""
            if caption:
                out += f'<p class="fig">{html.escape(caption)}</p>'
            return out
        return f'<p class="fig">{html.escape(item.text)}</p>'

    # one <p> per line: a bullet line needs a real indent for the DLA model to
    # read it as a list-item rather than run it into the paragraph above
    parts = [f'<p{" class=\"li\"" if _BULLET.match(line) else ""}>'
             f"{html.escape(line.strip())}</p>"
             for line in item.text.split("\n") if line.strip()]
    return "".join(parts)


def _page_html(page: LayoutPage) -> str:
    return "".join(_item_html(i) for i in page.items) or "<p></p>"


# The ToUnicode table MuPDF rebuilds maps a few glyphs to their "twin"
# characters — a completely different codepoint that draws identically, so
# nothing looks wrong on paper:
#
#   space     -> U+00A0 (non-breaking space)
#   hyphen    -> U+00AD (soft hyphen, an invisible character)
#   semicolon -> U+037E (Greek question mark)
#
# Every tool that reads the PDF back therefore receives a string without a
# single ordinary space or hyphen in it — tokenizers, BM25 and the word
# segmentation downstream all break with it.
_ALIASES = {0x00A0: 0x0020, 0x00AD: 0x002D, 0x037E: 0x003B}

_BFCHAR = re.compile(rb"^<([0-9a-fA-F]+)>\s*<([0-9a-fA-F]{4})>$")
_BFRANGE = re.compile(rb"^<([0-9a-fA-F]+)>\s*<([0-9a-fA-F]+)>\s*<([0-9a-fA-F]{4})>$")


def _alias_of(code: bytes) -> bytes | None:
    fixed = _ALIASES.get(int(code, 16))
    return f"{fixed:04x}".encode() if fixed is not None else None


def _fix_tounicode_aliases(doc: fitz.Document) -> int:
    """Point the mismapped glyphs back at the right characters.

    Returns the number of fonts that were fixed.
    """
    probes = [f"{code:04x}".encode() for code in _ALIASES]
    fixed = 0
    for xref in range(1, doc.xref_length()):
        if not doc.xref_is_stream(xref):
            continue
        try:
            data = doc.xref_stream(xref)
        except (RuntimeError, ValueError):
            continue
        if b"begincmap" not in data:
            continue
        lower = data.lower()
        if not any(probe in lower for probe in probes):
            continue

        out: list[bytes] = []
        section = b""
        changed = False
        for line in data.split(b"\n"):
            probe = line.strip()
            if probe.endswith(b"beginbfchar") or probe.endswith(b"beginbfrange"):
                section = b"char" if probe.endswith(b"beginbfchar") else b"range"
            elif probe in (b"endbfchar", b"endbfrange"):
                section = b""
            elif section == b"char":
                m = _BFCHAR.match(probe)
                alias = _alias_of(m.group(2)) if m else None
                if alias:
                    out.append(b"<" + m.group(1) + b"> <" + alias + b">")
                    changed = True
                    continue
            elif section == b"range":
                m = _BFRANGE.match(probe)
                alias = _alias_of(m.group(3)) if m else None
                if alias:
                    lo, hi, dst = m.group(1), m.group(2), m.group(3)
                    if lo.lower() == hi.lower():
                        out.append(b"<" + lo + b"> <" + hi + b"> <" + alias + b">")
                    else:
                        # a range that *starts* at the wrong character: split off
                        # the first code, and the rest still has to point at the
                        # next character to keep the order right
                        nxt = f"{int(lo, 16) + 1:04x}".encode()
                        after = f"{int(dst, 16) + 1:04x}".encode()
                        out.append(b"<" + lo + b"> <" + lo + b"> <" + alias + b">")
                        out.append(b"<" + nxt + b"> <" + hi + b"> <" + after + b">")
                    changed = True
                    continue
            out.append(line)

        if changed:
            doc.update_stream(xref, b"\n".join(out), compress=True)
            fixed += 1
    return fixed


class _Sheet:
    """The write cursor: the open page and the bottom edge of what is on it."""

    def __init__(self, writer, paper: fitz.Rect, frame: fitz.Rect,
                 first_page: int = 0):
        self.writer, self.paper, self.frame = writer, paper, frame
        self.device = None
        self.y = frame.y0
        self.page = first_page - 1     # no page opened yet

    @property
    def at_top(self) -> bool:
        return self.device is None or self.y <= self.frame.y0

    @property
    def room(self) -> float:
        return self.frame.y1 - self.y

    def avail(self) -> fitz.Rect:
        return fitz.Rect(self.frame.x0, self.y, self.frame.x1, self.frame.y1)

    def new_page(self) -> None:
        self.close()
        self.device = self.writer.begin_page(self.paper)
        self.y = self.frame.y0
        self.page += 1

    def close(self) -> None:
        if self.device is not None:
            self.writer.end_page()
            self.device = None


def _keep_together(items: list[PageItem]) -> list[tuple[list[PageItem], int]]:
    """Group blocks that must share a page, with the count of leading headings.

    A group = a run of consecutive heading lines plus the content block right
    after them. Consecutive headings (a parent section's heading followed by its
    first child's) have to travel together, and they must drag along at least a
    fragment of their own content — otherwise the heading is left stranded at the
    foot of the page, which is the very fault being fixed here.
    """
    groups: list[list[PageItem]] = []
    heads: list[int] = []
    for item in items:
        if item.kind == "heading":
            if groups and len(groups[-1]) == heads[-1]:
                groups[-1].append(item)      # extend the open run of headings
                heads[-1] += 1
                continue
            groups.append([item])
            heads.append(1)
        elif groups and len(groups[-1]) == heads[-1] and heads[-1]:
            groups[-1].append(item)          # the run's first piece of content
        else:
            groups.append([item])
            heads.append(0)
    return list(zip(groups, heads))


def _natural_height(group: list[PageItem]) -> float:
    """The height this group *needs* in order not to be squashed — figures only."""
    return max((float(i.meta.get("height") or 0)
                for i in group if i.kind == "figure"), default=0.0)


def _draw_flow(sheet: _Sheet, items: list[PageItem], css: str,
               archive: fitz.Archive,
               forced: set[int]) -> list[tuple[int, int, bool]]:
    """Pour the blocks onto pages.

    Returns a placement trace: (block index, starting page, whether the block
    opens with a heading). `forced` holds the blocks that must start on a fresh
    page — the findings of the previous rebuild pass.
    """
    trace: list[tuple[int, int, bool]] = []
    sheet.new_page()
    for index, (group, heads) in enumerate(_keep_together(items)):
        markup = "".join(_item_html(i) for i in group)
        if not markup:
            continue
        # When the remaining gap is too short, MuPDF *shrinks the image* to fit
        # rather than pushing it to the next page — a 630pt flowchart landing at
        # the foot of a page is squeezed to under a third of its size and the
        # text in its boxes becomes unreadable. Opening a new page before placing
        # it is the only way to keep a figure at full size.
        need = min(_natural_height(group), sheet.frame.height)
        if not sheet.at_top and (index in forced
                                 or sheet.room < max(MIN_PLACE_PT, need)):
            sheet.new_page()

        story = fitz.Story(html=markup, user_css=css, archive=archive)
        more, filled = story.place(sheet.avail())
        trace.append((index, sheet.page, bool(heads)))

        for _ in range(_MAX_PAGES_PER_BLOCK):
            story.draw(sheet.device)
            sheet.y = fitz.Rect(filled).y1
            if not more:
                break
            sheet.new_page()
            more, filled = story.place(sheet.avail())
    sheet.close()
    return trace


def _stranded_pages(doc: fitz.Document) -> set[int]:
    """Pages that end on a heading line with no content below it.

    MuPDF decides page breaks after laying out the text, and `place()` cannot
    report what actually appeared: given a table row taller than the remaining
    gap, it reports the space as used while in fact drawing nothing but the
    heading line. The only reliable check is to read the rendered page back and
    see whether anything follows the heading.
    """
    stranded: set[int] = set()
    for number in range(doc.page_count - 1):
        page = doc[number]
        bottom = 0.0
        last_size = 0.0
        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                size = max((s["size"] for s in line["spans"] if s["text"].strip()),
                           default=0.0)
                if not size:
                    continue
                if line["bbox"][3] > bottom:
                    bottom, last_size = line["bbox"][3], size
        if last_size <= BODY_PT + 0.4:
            continue
        if any(info["bbox"][3] > bottom for info in page.get_image_info()):
            continue                    # a figure follows the heading, not blank space
        stranded.add(number)
    return stranded


def _render(pages: list[LayoutPage], css: str, archive: fitz.Archive,
            paper: fitz.Rect, frame: fitz.Rect,
            forced: dict[int, set[int]]) -> tuple[bytes, dict[int, tuple[int, int]]]:
    """Render the whole document.

    Returns (PDF bytes, page -> the last block placed on that page).
    """
    buffer = io.BytesIO()
    writer = fitz.DocumentWriter(buffer)
    last_on_page: dict[int, tuple[int, int]] = {}
    next_page = 0
    for index, page in enumerate(pages):
        # Every LayoutPage opens a new sheet: in one-page-per-section mode that
        # is exactly the page boundary to preserve.
        sheet = _Sheet(writer, paper, frame, next_page)
        for group, page_no, is_head in _draw_flow(sheet, page.items, css,
                                                  archive,
                                                  forced.get(index, set())):
            if is_head:
                last_on_page[page_no] = (index, group)
            else:
                last_on_page.pop(page_no, None)
        next_page = sheet.page + 1
    writer.close()
    return buffer.getvalue(), last_on_page


def write_pdf(pages: list[LayoutPage], dst: str, figure_dir: str | None = None) -> None:
    font = _font_files()
    archive = fitz.Archive()
    if font:
        archive.add(font[0])
    if figure_dir and os.path.isdir(figure_dir):
        archive.add(figure_dir)

    css = _css(font)
    paper = fitz.paper_rect("a4")
    frame = paper + (57, 57, -57, -57)      # 2cm margins

    os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)
    # Render in memory before writing: the raw build embeds the entire system
    # font (several MB), which has to be subset before saving.
    #
    # Several rebuild passes: each one reads the pages just rendered, finds the
    # heading lines left stranded at the foot of a page and marks them so the
    # next pass opens a fresh page right before them. The set of marked blocks
    # only grows, so the loop is guaranteed to terminate.
    forced: dict[int, set[int]] = {}
    data, last_on_page = _render(pages, css, archive, paper, frame, forced)
    for _ in range(_MAX_REFLOW_PASSES):
        doc = fitz.open("pdf", data)
        try:
            stranded = _stranded_pages(doc)
        finally:
            doc.close()
        fresh = False
        for page_no in stranded:
            owner = last_on_page.get(page_no)
            if owner is None:
                continue
            index, group = owner
            if group not in forced.setdefault(index, set()):
                forced[index].add(group)
                fresh = True
        if not fresh:
            break
        data, last_on_page = _render(pages, css, archive, paper, frame, forced)

    doc = fitz.open("pdf", data)
    try:
        doc.subset_fonts()          # keep only the characters actually used
    except (AttributeError, RuntimeError):
        pass
    _fix_tounicode_aliases(doc)     # after subsetting, which rebuilds ToUnicode
    try:
        doc.save(dst, garbage=4, deflate=True)
    finally:
        doc.close()


def write(pages: list[LayoutPage], dst: str, figure_dir: str | None = None) -> None:
    """Write the file, choosing the writer from `dst`'s extension."""
    ext = os.path.splitext(dst)[1].lower()
    if ext == ".pdf":
        write_pdf(pages, dst, figure_dir=figure_dir)
    elif ext == ".docx":
        write_docx(pages, dst)
    else:
        raise ValueError(f"Unsupported format: {ext}")
