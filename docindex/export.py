"""Write out the cleaned document.

There are two ways to do it:

* **Rebuild the layout** (default) — content flows from the extracted outline,
  every section name alone on its line at a size that shrinks with depth. The
  output format can be `.docx` or `.pdf`.
* **Keep the original layout** — PDF in, PDF out; DOCX in, DOCX out; only the
  items selected in `CleanOptions` are stripped: logos, cover art, headers and
  footers, table of contents.
"""
from __future__ import annotations

import math
import os
import re
from dataclasses import dataclass

import docx
import fitz
from docx.oxml.ns import qn

from . import render
from .extract_pdf import (
    _DOT_LEADER, _PAGE_NUM, _TOC_TITLE, FOOTER_ZONE, HEADER_ZONE,
    _collect_repeats, _line_text, _norm_key, _page_dict, is_boilerplate,
)
from .chunker import MAX_TOKENS
from .extract_docx import _classify_docx_image, _paragraph_images
from .images import collect_pdf_images
from .layout import LINES_PER_PAGE, build_pages
from .models import Section, document_title

# After the noise is stripped, a page with fewer characters than this is blank
EMPTY_PAGE_CHARS = 15

# Suffix of the normalised copy. The name alone tells the original document
# apart from the RAG build, even when both sit in the same folder.
FORMALIZED_SUFFIX = "_formalized"


@dataclass
class CleanOptions:
    """Pick individually what to strip while keeping the original layout.

    Not everyone wants everything stripped. Some only need the logo and cover
    art gone so the file gets lighter while the text stays exactly where it was;
    removing headers, footers and the table of contents changes the content, so
    the user has to ask for it deliberately.
    """

    drop_logo: bool = True            # logos, repeated ornaments, page backgrounds
    drop_cover: bool = True           # large cover art on the first page
    drop_header_footer: bool = True   # header / footer text
    drop_toc: bool = True             # the table of contents (and TOC-only pages)

    @property
    def touches_text(self) -> bool:
        return self.drop_header_footer or self.drop_toc


def _noise_rects(page: fitz.Page, repeats: set[str],
                 drop_header_footer: bool = True,
                 drop_toc: bool = True) -> tuple[list[fitz.Rect], bool]:
    """Find the text areas to erase on a page.

    Returns (list of rectangles, whether the page is nothing but a table of
    contents).
    """
    height = page.rect.height
    rects: list[fitz.Rect] = []
    leaders: list[float] = []
    toc_title_y: float | None = None
    kept_before_toc = 0

    lines = []
    for block in _page_dict(page)["blocks"]:
        if block["type"] != 0:
            continue
        for line in block["lines"]:
            text = _line_text(line)
            if text:
                lines.append((text, fitz.Rect(line["bbox"])))

    for text, rect in lines:
        if _DOT_LEADER.search(text):
            leaders.append(rect.y0)
        elif _TOC_TITLE.match(text):
            toc_title_y = rect.y0

    toc_start: float | None = None
    if len(leaders) >= 4:
        toc_start = min(leaders)
        if toc_title_y is not None and toc_title_y <= toc_start:
            toc_start = toc_title_y

    for text, rect in lines:
        ratio = rect.y0 / height if height else 0
        in_margin = ratio <= HEADER_ZONE or ratio >= FOOTER_ZONE
        noise = is_boilerplate(text, ratio, repeats)
        if not noise and in_margin and len(text) < 90:
            tail = text.split()[-1] if text.split() else ""
            if _PAGE_NUM.search(tail):
                stripped = _norm_key(re.sub(r"\s*\d+\s*/\s*\d+\s*$", "", text))
                noise = stripped in repeats

        in_toc = toc_start is not None and rect.y0 >= toc_start
        is_toc = in_toc or bool(_DOT_LEADER.search(text))
        if (noise and drop_header_footer) or (is_toc and drop_toc):
            rects.append(rect)
        # Body text is counted by classification, not by the user's choice:
        # turning TOC removal off must not turn a TOC page into a page with
        # real content.
        if not noise and not is_toc and not in_margin:
            kept_before_toc += len(text)

    only_toc = drop_toc and toc_start is not None and kept_before_toc < 40
    return rects, only_toc


def _has_content(page: fitz.Page) -> bool:
    """Is anything worth keeping left on the page — text, images or strokes?

    PyMuPDF's `delete_image` does not actually remove an image; it replaces it
    with a transparent 1x1 cell, so a removed image is still listed on the page.
    Only the real dimensions say whether anything still shows up on paper.
    """
    if len(page.get_text("text").strip()) >= EMPTY_PAGE_CHARS:
        return True
    if any(info[2] > 1 and info[3] > 1 for info in page.get_images(full=True)):
        return True
    return bool(page.get_drawings())


def clean_pdf(src: str, dst: str, opts: CleanOptions | None = None) -> dict:
    """Write the cleaned PDF. Returns statistics on what was stripped."""
    opts = opts or CleanOptions()
    doc = fitz.open(src)
    images = collect_pdf_images(doc, treat_first_page_as_cover=opts.drop_cover)
    # Only scan for repeated text when text is actually touched — a full-document
    # scan is not cheap
    repeats = _collect_repeats(doc) if opts.touches_text else set()
    drop_kinds = {k for k, on in (("logo", opts.drop_logo),
                                  ("cover", opts.drop_cover)) if on}

    removed_images = 0
    removed_text_zones = 0
    kept_figures = 0
    blank_pages: list[int] = []

    for pno in range(doc.page_count):
        page = doc[pno]

        # 1) Strip logos, ornaments and cover art.
        #    This must go by each image's own xref. Area-based redaction erases
        #    every image *intersecting* the area — and logos often sit on top of
        #    a large figure, so that approach would take the content with them.
        drop = [i for i in images.get(pno + 1, []) if i.kind in drop_kinds]
        kept_figures += sum(1 for i in images.get(pno + 1, []) if i.kind == "figure")
        leftover: list[fitz.Rect] = []
        for img in drop:
            if img.xref:
                try:
                    page.delete_image(img.xref)
                    removed_images += 1
                    continue
                except (ValueError, RuntimeError):
                    pass
            rect = fitz.Rect(img.bbox) & page.rect
            if not rect.is_empty:
                leftover.append(rect)
        if leftover:
            # images inlined in the content stream have no xref to remove
            for rect in leftover:
                page.add_redact_annot(rect)
            page.apply_redactions(images=1, graphics=0, text=0)
            removed_images += len(leftover)

        # 2) A cover page is usually empty once stripped — it was never more than
        #    one image. Left in place, the cleaned file opens on a blank sheet,
        #    and a DLA model can do nothing with an empty page either.
        if drop and not _has_content(page):
            blank_pages.append(pno)
            continue

        # 3) Strip header/footer text and the table of contents, leaving figures
        if not opts.touches_text:
            continue
        rects, only_toc = _noise_rects(page, repeats,
                                       drop_header_footer=opts.drop_header_footer,
                                       drop_toc=opts.drop_toc)
        if rects:
            for rect in rects:
                page.add_redact_annot(rect)
            page.apply_redactions(images=0, graphics=0, text=0)
            removed_text_zones += len(rects)

        if only_toc and len(page.get_text("text").strip()) < EMPTY_PAGE_CHARS:
            blank_pages.append(pno)

    # 4) Drop pages that hold nothing but a TOC, and covers just stripped bare
    for pno in sorted(set(blank_pages), reverse=True):
        doc.delete_page(pno)

    os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)
    # use_objstms packs small objects into compressed streams. Without it the
    # written file nearly doubles the original size, even with content removed.
    #
    # garbage=1 rather than 4: level 4 makes MuPDF hash the content of *every*
    # stream looking for duplicates, and on an image-heavy document that costs
    # minutes rather than seconds — 86s against 0.1s on one real file. Expensive
    # and pointless here: the noise images were already removed by xref above, so
    # level 1 (drop objects nothing references any more) produced a file *exactly*
    # the size of level 4 on all three of the heaviest documents (2.77 MB,
    # 6.64 MB, 0.94 MB — identical to the last digit).
    try:
        doc.save(dst, garbage=1, deflate=True, use_objstms=1)
    except TypeError:  # older PyMuPDF without this parameter
        doc.save(dst, garbage=1, deflate=True)
    pages_out = doc.page_count
    doc.close()

    return {
        "images_removed": removed_images,
        "figures_kept": kept_figures,
        "text_zones_removed": removed_text_zones,
        "pages_removed": len(blank_pages),
        "pages_out": pages_out,
    }


def _has_image(element) -> bool:
    """Does this run hold an image (both modern and legacy Word embeddings)?"""
    return bool(element.findall(qn("w:drawing")) or element.findall(qn("w:pict")))


def _clear_headers_footers(document, drop_text: bool = True,
                           drop_images: bool = True) -> int:
    """Clear the header and footer content of every section.

    Text and images can be removed independently: the company logo is almost
    always in the header, so someone who only wants the logo gone still has to
    touch the header without losing the "issued together with Decision no…"
    line that also lives there.
    """
    removed = 0
    for section in document.sections:
        for part in (section.header, section.footer,
                     section.even_page_header, section.even_page_footer,
                     section.first_page_header, section.first_page_footer):
            if part is None:
                continue
            for para in part.paragraphs:
                if drop_text and drop_images:
                    if para.text.strip() or _has_image(para._p):
                        removed += 1
                    for child in list(para._p):
                        if child.tag != qn("w:pPr"):
                            para._p.remove(child)
                    continue
                for run in list(para.runs):
                    has_image = _has_image(run._r)
                    if (has_image and drop_images) or (not has_image and drop_text):
                        run._r.getparent().remove(run._r)
                        removed += 1
            if not drop_text:
                continue
            for table in part.tables:
                table._tbl.getparent().remove(table._tbl)
                removed += 1
    return removed


def clean_docx(src: str, dst: str, opts: CleanOptions | None = None) -> dict:
    """Write the cleaned DOCX, with the content figures still in the file."""
    opts = opts or CleanOptions()
    document = docx.Document(src)

    # Word has no notion of a "cover page" while editing, so cover art and logos
    # in a .docx both fall to the same size-based classification.
    drop_images = opts.drop_logo or opts.drop_cover
    removed_headers = 0
    if opts.drop_header_footer or drop_images:
        removed_headers = _clear_headers_footers(
            document, drop_text=opts.drop_header_footer, drop_images=drop_images)
    removed_images = 0
    kept_figures = 0

    for para in document.paragraphs:
        for run in list(para.runs):
            if not run._r.findall(qn("w:drawing")):
                continue
            infos = _paragraph_images(para)
            # A small or unmeasurable image is a logo/ornament -> drop the run
            keep = not drop_images or any(
                _classify_docx_image(i) == "figure" for i in infos)
            if keep:
                kept_figures += 1
                continue
            run._r.getparent().remove(run._r)
            removed_images += 1

    os.makedirs(os.path.dirname(os.path.abspath(dst)), exist_ok=True)
    document.save(dst)

    return {
        "images_removed": removed_images,
        "figures_kept": kept_figures,
        "header_footer_parts_cleared": removed_headers,
        "text_zones_removed": 0,
        "pages_removed": 0,
    }


def out_dir_for(src: str, in_root: str, out_dir: str) -> str:
    """Output folder for one file, mirroring the input folder tree.

    Scanning a whole tree, two subfolders may well hold two files with the same
    name (an original and a table-flattened copy, say). Dumped into one place,
    the later file overwrites the earlier one and its results are lost without a
    word.
    """
    rel = os.path.relpath(os.path.dirname(os.path.abspath(src)),
                          os.path.abspath(in_root))
    return out_dir if rel in (".", "") else os.path.join(out_dir, rel)


def _stem_with_suffix(name: str, suffix: str) -> str:
    """Append the suffix to a file name, without repeating one already there.

    Re-running the tool on its own output folder is routine; without this step
    names grow into `quy-dinh_formalized_formalized.pdf`.
    """
    return name if not suffix or name.endswith(suffix) else f"{name}{suffix}"


def out_path(src: str, out_dir: str, out_format: str = "same",
             suffix: str = FORMALIZED_SUFFIX) -> str:
    """Path of the output file; `out_format` is "same", "docx" or "pdf"."""
    name, ext = os.path.splitext(os.path.basename(src))
    ext = ext.lower() if out_format in ("same", "", None) else f".{out_format.lower()}"
    dst = os.path.join(out_dir, f"{_stem_with_suffix(name, suffix)}{ext}")
    _refuse_to_overwrite_source(src, dst)
    return dst


def _refuse_to_overwrite_source(src: str, dst: str) -> None:
    """Refuse the case where the output file is exactly the input file.

    The `_formalized` suffix already separates the rebuild from the original,
    but running the tool on an already-normalised file does not double the
    suffix, so the output name matches the source name exactly — and overwriting
    it destroys the document beyond recovery.
    """
    if os.path.abspath(src) == os.path.abspath(dst):
        raise ValueError(
            f"the output file is the source file ({os.path.basename(src)}) — "
            "choose an -o folder other than the one holding the input document"
        )


def _pages_written(dst: str, pages, lines_per_page: int) -> int:
    """Page count of the file just written.

    In continuous-flow mode the renderer decides the page breaks, so the count
    comes from the real file. With .docx, Word paginates when the file is opened
    and the count cannot be read back — so it is estimated from the line count.
    """
    if os.path.splitext(dst)[1].lower() == ".pdf":
        try:
            with fitz.open(dst) as doc:
                return doc.page_count
        except (RuntimeError, ValueError):
            pass
    if len(pages) > 1:
        return len(pages)
    total = sum(p.lines for p in pages)
    return max(1, math.ceil(total / lines_per_page)) if total else 0


def rebuild_document(sections: list[Section], src: str, out_dir: str,
                     out_format: str = "pdf", suffix: str = FORMALIZED_SUFFIX,
                     figure_dir: str | None = None, drop_cover: bool = True,
                     lines_per_page: int = LINES_PER_PAGE,
                     max_tokens: int = MAX_TOKENS,
                     page_per_section: bool = False,
                     doc_title: str | None = None) -> tuple[str, dict]:
    """Rebuild the document with the outline tree laid bare for a RAG stack.

    The content comes from the extracted outline, so logos, headers, footers and
    the table of contents were dropped upstream — nothing to strip again here.

    The default output is `.pdf`: a DLA model reads PDF far more reliably than
    `.docx`, so that is the format to feed the RAG stack. Content flows
    continuously like an ordinary document; `page_per_section=True` returns to
    one page per section, and `max_tokens` then caps each page.
    """
    dst = out_path(src, out_dir, out_format, suffix)
    # A large image on a PDF's first page is nearly always cover art. Word has no
    # notion of pages while editing, so the rule is not applied to .docx.
    cover = drop_cover and os.path.splitext(src)[1].lower() == ".pdf"
    title = document_title(src) if doc_title is None else doc_title
    pages = build_pages(sections, lines_per_page=lines_per_page,
                        max_tokens=max_tokens, drop_cover=cover,
                        page_per_section=page_per_section, doc_title=title)
    render.write(pages, dst, figure_dir=figure_dir)

    figures = sum(1 for p in pages for i in p.items if i.kind == "figure")
    return dst, {
        "images_removed": 0,
        "figures_kept": figures,
        "text_zones_removed": 0,
        "pages_removed": 0,
        "pages_out": _pages_written(dst, pages, lines_per_page),
        # count the real sections built rather than pages: continuous-flow mode
        # has exactly one "layout page"
        "sections_out": sum(1 for s in sections if not s.is_preamble),
        "doc_title": title,
        "layout": "per-section" if page_per_section else "flow",
    }


def clean_document(src: str, out_dir: str, suffix: str = FORMALIZED_SUFFIX,
                   opts: CleanOptions | None = None) -> tuple[str, dict]:
    """Clean one document, keeping the input format and layout."""
    name, ext = os.path.splitext(os.path.basename(src))
    ext = ext.lower()
    dst = os.path.join(out_dir, f"{_stem_with_suffix(name, suffix)}{ext}")
    _refuse_to_overwrite_source(src, dst)

    if ext == ".pdf":
        stats = clean_pdf(src, dst, opts)
    elif ext == ".docx":
        stats = clean_docx(src, dst, opts)
    else:
        raise ValueError(f"Unsupported format: {ext}")
    return dst, stats
