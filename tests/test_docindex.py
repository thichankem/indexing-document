"""Check the tool's core constraints against a corpus of real documents.

The fixture text is Vietnamese because that is the language of the documents
this tool processes; test names, docstrings and assertion messages are English.
"""
from __future__ import annotations

import os
import re
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docindex.chunker import (
    ChunkConfig, _split_long_text, _split_table, est_tokens,
)
from docindex.headings import (
    _match_heading, build_sections, detect_headings, shorten_title,
)
from docindex.images import CAPTION, DocImage
from docindex.models import (
    PREAMBLE_TITLE, Block, Section, clean_text, normalize_symbols,
    rows_to_markdown,
)
from docindex.numbering import _format_counter, _to_letter, _to_roman
from docindex.pipeline import iter_input_files, process_file

TEST_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "testing file")
CFG = ChunkConfig()


def _docs():
    if not os.path.isdir(TEST_DIR):
        pytest.skip("the 'testing file' folder does not exist")
    files = iter_input_files(TEST_DIR)
    if not files:
        pytest.skip("no pdf/docx files available to test against")
    return files


@pytest.fixture(scope="module")
def processed():
    out = []
    for path in _docs():
        stats: dict = {}
        chunks, sections, src = process_file(path, CFG, stats=stats)
        out.append((path, chunks, sections, src, stats))
    return out


# --- unit level ---------------------------------------------------------

def test_roman_and_letter_counters():
    assert _to_roman(4) == "iv"
    assert _to_roman(2026) == "mmxxvi"
    assert _to_letter(1) == "a"
    assert _to_letter(27) == "aa"
    assert _format_counter(3, "upperRoman") == "III"
    assert _format_counter(2, "decimalZero") == "02"


def test_clean_text_removes_invisible_chars():
    assert clean_text("1.​Phạm vi") == "1.Phạm vi"
    assert clean_text("a  b") == "a b"


def test_heading_patterns():
    assert _match_heading("ĐIỀU 2: MỘT SỐ QUY ĐỊNH")[0] == "ĐIỀU 2"
    assert _match_heading("2.3.1. Nghĩa vụ cung cấp")[0] == "2.3.1"
    assert _match_heading("PHẦN 1: QUY TẮC CHUNG")[3] == "part"
    assert _match_heading("Đây là câu văn bình thường") is None


def test_reference_phrase_is_not_a_heading():
    """A cross-reference fragment once matched and swallowed every section after it."""
    blocks = [
        Block(text="Phần 1 này", page=1, bold=False),
        Block(text="ĐIỀU 3: PHÍ BẢO HIỂM", page=1, bold=True),
    ]
    detect_headings(blocks, "pdf")
    assert blocks[0].kind != "heading"
    assert blocks[1].kind == "heading"


def test_split_long_text_respects_limit():
    text = ". ".join(f"Câu số {i} trong đoạn văn dài" for i in range(80))
    parts = _split_long_text(text, 200)
    assert parts and all(len(p) <= 200 for p in parts)


def test_split_table_repeats_header():
    md = "\n".join(["| A | B |", "| --- | --- |"] + [f"| {i} | giá trị {i} |" for i in range(60)])
    parts = _split_table(md, 300)
    assert len(parts) > 1
    assert all(p.startswith("| A | B |") for p in parts)


# --- constraints on real documents --------------------------------------

def test_every_document_produces_chunks(processed):
    for path, chunks, _sections, _src, _st in processed:
        assert chunks, f"no chunks produced for {os.path.basename(path)}"


def test_chunk_never_spans_two_pages(processed):
    """The core constraint: everything in one chunk comes from a single page.

    Each line of a chunk is matched against the page it came from in the
    document. A line that occurs on exactly one page, but not the chunk's page,
    is a bug.
    """
    for path, chunks, sections, _src, _st in processed:
        pages_of_line: dict[str, set[int]] = {}
        for section in sections:
            for block in section.blocks:
                for line in block.text.split("\n"):
                    line = line.strip()
                    if len(line) > 40:
                        pages_of_line.setdefault(line, set()).add(block.page)

        for chunk in chunks:
            assert isinstance(chunk.page, int) and chunk.page >= 1
            for line in chunk.raw_text.split("\n"):
                line = line.strip()
                pages = pages_of_line.get(line)
                if pages and len(pages) == 1:
                    assert chunk.page in pages, (
                        f"{os.path.basename(path)} {chunk.chunk_id}: line belongs to "
                        f"page {pages} but the chunk declares page {chunk.page}"
                    )


def test_chunk_never_exceeds_token_limit(processed):
    """The RAG stack's hard limit: 512 tokens. Over it, embedding truncates."""
    for path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            assert chunk.est_tokens <= CFG.max_tokens, (
                f"{os.path.basename(path)} {chunk.chunk_id}: {chunk.est_tokens} "
                f"token > {CFG.max_tokens}"
            )
            assert chunk.est_tokens == est_tokens(chunk.text)


def test_token_estimate_is_not_optimistic():
    """The token estimate must run high, never low.

    Undercount and a chunk passes every check here, only to be truncated at
    embedding time — a silent failure, so overestimating is the safe side.
    """
    plain = "Khách hàng cá nhân ưu tiên được hưởng chính sách lãi suất ưu đãi."
    assert est_tokens(plain) >= len(plain.split())
    table = "| Phân hạng | Mức giảm |\n| --- | --- |\n| Diamond | 0,2% |"
    assert est_tokens(table) >= table.count("|")
    assert est_tokens("") == 0


def test_chunk_ids_unique(processed):
    for _path, chunks, _sections, _src, _st in processed:
        ids = [c.chunk_id for c in chunks]
        assert len(ids) == len(set(ids))


def test_chunk_links_are_consistent(processed):
    for _path, chunks, _sections, _src, _st in processed:
        by_id = {c.chunk_id: c for c in chunks}
        for chunk in chunks:
            if chunk.prev_chunk_id:
                assert by_id[chunk.prev_chunk_id].next_chunk_id == chunk.chunk_id
            if chunk.next_chunk_id:
                assert by_id[chunk.next_chunk_id].prev_chunk_id == chunk.chunk_id


def test_continuation_flags_match_section(processed):
    """is_continued holds only when the next chunk belongs to the same section."""
    for _path, chunks, _sections, _src, _st in processed:
        by_id = {c.chunk_id: c for c in chunks}
        for chunk in chunks:
            if chunk.is_continued:
                nxt = by_id[chunk.next_chunk_id]
                assert nxt.section_path == chunk.section_path
                assert nxt.is_continuation
            if chunk.is_continuation:
                prev = by_id[chunk.prev_chunk_id]
                assert prev.section_path == chunk.section_path


def test_section_path_starts_the_text(processed):
    for _path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            if chunk.section_path:
                assert chunk.text.startswith(chunk.section_path.split(" > ")[0][:20])


def test_most_chunks_belong_to_a_section(processed):
    """When heading detection breaks, most chunks fall into the preamble."""
    for path, chunks, _sections, _src, _st in processed:
        orphan = [c for c in chunks if c.section_path == PREAMBLE_TITLE]
        assert len(orphan) <= max(2, len(chunks) * 0.15), os.path.basename(path)


def test_headings_are_detected(processed):
    for path, _chunks, sections, _src, _st in processed:
        numbered = [s for s in sections if s.number]
        assert numbered, f"no headings detected in {os.path.basename(path)}"


def test_no_dot_leader_garbage(processed):
    """TOC pages must be dropped, never yielding a chunk full of dot leaders."""
    for _path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            assert "........" not in chunk.text


# --- dropping logos / keeping figures ------------------------------------

def test_caption_pattern():
    assert CAPTION.match("Hình 1: Quy trình xử lý")
    assert CAPTION.match("Biểu đồ 2 - Tăng trưởng")
    assert CAPTION.match("Sơ đồ 3. Cơ cấu tổ chức")
    assert not CAPTION.match("Hình thức thanh toán theo quy định")


def test_figure_placeholder_format():
    img = DocImage(page=2, bbox=(0, 0, 300, 200), width=300, height=200,
                   kind="figure", reason="test", caption="Biểu đồ 1: Doanh thu")
    text = img.placeholder()
    assert text.startswith("[FIGURE:")
    assert "Biểu đồ 1" in text and "300x200" in text


def test_logos_are_dropped_and_not_in_text(processed):
    """A logo repeated in the header or footer must never reach a chunk."""
    total_logos = sum(st.get("logos_dropped", 0) for _p, _c, _s, _src, st in processed)
    assert total_logos > 0, "no logos were dropped across the real corpus"

    for _path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            for line in chunk.raw_text.split("\n"):
                if line.strip().startswith("[FIGURE:"):
                    continue
                assert ".png" not in line.lower() and ".jpeg" not in line.lower()


def test_figures_are_kept_with_metadata(processed):
    """Content figures keep a placeholder plus metadata, and are never dropped."""
    with_fig = [c for _p, chunks, _s, _src, _st in processed for c in chunks if c.has_figure]
    assert with_fig, "no figures were kept"
    for chunk in with_fig:
        assert chunk.figures
        assert "[FIGURE:" in chunk.raw_text
        for fig in chunk.figures:
            assert fig["width"] and fig["height"]


def test_figure_flag_matches_content(processed):
    for _path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            if "[FIGURE:" in chunk.raw_text:
                assert chunk.has_figure, f"{chunk.chunk_id} holds a figure but is not flagged"


def test_letter_marker_nests_under_decimal():
    """"a)" is a child of "2.", not its sibling.

    Real Word documents routinely declare an "a) b) c)" list at the same `ilvl`
    as "1. 2. 3.". Trust that level and the outline flattens, costing every
    chunk below it a tier of its title.
    """
    from docindex.headings import build_sections

    blocks = [
        Block(text="Chính sách lãi suất huy động", page=1, bold=True,
              number="2.", level=1, meta={"list_format": "decimal"}),
        Block(text="Phạm vi áp dụng", page=1, bold=True,
              number="a)", level=1, meta={"list_format": "lowerLetter"}),
    ]
    detect_headings(blocks, "docx")
    sections = build_sections(blocks)
    assert [s.level for s in sections] == [1, 2]
    assert sections[1].path == ["2. Chính sách lãi suất huy động", "a) Phạm vi áp dụng"]


def test_sentence_starting_with_a_number_is_not_a_heading():
    """"30 (ba mươi) ngày tuổi…" is prose wrapped by the PDF, not section 30.

    Mistaking a line like that does more than spawn a phantom section: every
    real section after it drops down to become its child, and the titles of the
    whole branch go wrong with it.
    """
    from docindex.headings import build_sections

    blocks = [
        Block(text="1.25 Người được bảo hiểm: là cá nhân được chấp thuận",
              page=1, bold=True),
        Block(text="30 (ba mươi) ngày tuổi đến 70 (bảy mươi) tuổi tại thời điểm",
              page=1, bold=True),
        Block(text="1.26 Người thụ hưởng: là cá nhân được chỉ định", page=1, bold=True),
    ]
    detect_headings(blocks, "pdf")
    assert blocks[1].kind != "heading"
    sections = build_sections(blocks)
    assert [s.number for s in sections] == ["1.25", "1.26"]
    assert [s.level for s in sections] == [1, 1], "a real section was demoted"


def test_padded_number_is_a_quantity_not_an_index():
    """"04 (bốn) Năm hợp đồng" is a quantity — section numbers are never zero-padded."""
    blocks = [
        Block(text="3.1 Quy định chung về Phí bảo hiểm", page=1, bold=True),
        Block(text="04 (bốn) Năm hợp đồng đầu tiên.", page=1, bold=True),
        Block(text="3.2 Thời gian gia hạn đóng phí", page=1, bold=True),
    ]
    detect_headings(blocks, "pdf")
    assert blocks[1].kind != "heading"
    assert [b.number for b in blocks if b.kind == "heading"] == ["3.1", "3.2"]


def test_missing_parent_heading_does_not_drop_the_subtree():
    """When extraction misses a parent heading, its children must still survive."""
    blocks = [
        Block(text="2.2.5 Nội dung mục con cuối", page=1, bold=True),
        Block(text="2.3.1 Nội dung mục con đầu của nhánh sau", page=1, bold=True),
    ]
    detect_headings(blocks, "pdf")
    assert [b.number for b in blocks if b.kind == "heading"] == ["2.2.5", "2.3.1"]


def test_section_name_stops_at_the_colon():
    """The section name ends at the colon, even when the line is under the ceiling.

    Administrative documents run whole sentences into the heading line. Left
    alone, the sentence is set in bold at heading size and every chunk title in
    the section reads like prose.
    """
    assert shorten_title(
        "Hợp đồng bảo hiểm: là tất cả văn bản thể hiện sự thỏa thuận giữa hai bên"
    ) == "Hợp đồng bảo hiểm"
    assert shorten_title("Điều kiện áp dụng:") == "Điều kiện áp dụng"
    # a colon mid-way through a long sentence is not a name boundary
    long_head = "Trường hợp Khách hàng " + "rất dài " * 20
    assert shorten_title(long_head + ": nội dung") != long_head.strip()
    # do not cut on a clock time or a ratio
    assert shorten_title("Khung 8:30 tới 17:00") == "Khung 8:30 tới 17:00"


def test_run_on_heading_moves_its_body_down(processed):
    """Text run into the heading line belongs in the body, not on the heading."""
    from docindex.layout import _content_items, _heading_and_rest

    checked = 0
    for _path, _chunks, sections, _src, _st in processed:
        for section in sections:
            # A section with merged numbers carries several names joined, so it
            # runs longer than its own original heading line — by design, not a bug
            if ":" not in section.full_heading or not section.number or section.is_merged:
                continue
            name, rest = _heading_and_rest(section)
            if not rest:
                continue
            checked += 1
            assert len(name) <= len(section.full_heading)
            assert ":" not in name.rstrip(":"), f"content still stuck to the name: {name!r}"
            body = " ".join(i.text for i in _content_items(section, True))
            assert rest.split()[0] in body, "the run-on text fell out of the body"
    assert checked, "no section runs content into its heading, nothing to check"


def test_hyphenated_word_split_across_lines_is_rejoined():
    """"…và Dai-" at the end of a line joins "ichi Life…" into "Dai-ichi Life"."""
    from docindex.extract_pdf import _merge_wrapped

    lines = [
        Block(text="Hợp đồng bảo hiểm là thỏa thuận giữa Bên mua và Dai-", page=1,
              bold=True, size=11, meta={"x0": 70, "y": 100}),
        Block(text="ichi Life Việt Nam trên cơ sở yêu cầu bảo hiểm.", page=1,
              bold=False, size=11, meta={"x0": 70, "y": 114}),
    ]
    merged = _merge_wrapped(lines)
    assert len(merged) == 1
    assert "Dai-ichi Life Việt Nam" in merged[0].text


def test_short_sections_are_merged_without_losing_text():
    """Merging a short section loses its node in the tree, but never its words."""
    from docindex.headings import build_sections, merge_short_sections

    blocks = [
        Block(text="1. Giải thích từ ngữ", page=1, bold=True),
        Block(text="1.1 Bệnh: là tình trạng sức khỏe kém", page=1, bold=True),
        Block(text="1.2 Khoản nợ: là khoản tiền đến hạn", page=1, bold=True),
    ]
    detect_headings(blocks, "pdf")
    sections = build_sections(blocks)
    assert len(sections) == 3

    merged = merge_short_sections(sections, min_tokens=60, max_tokens=512)
    assert len(merged) == 1, "the short sections were not merged"
    text = merged[0].heading_text + " " + " ".join(b.text for b in merged[0].blocks)
    for word in ("Bệnh", "sức khỏe kém", "Khoản nợ", "đến hạn"):
        assert word in text, f"word lost during the merge: {word}"

    kept = merge_short_sections(sections, min_tokens=0, max_tokens=512)
    assert len(kept) == 3, "min_tokens=0 must merge nothing"


def test_upper_roman_ranks_above_decimal():
    """"I. THÔNG TIN CHUNG" is a high level, ranking above "10. Điều kiện"."""
    assert _match_heading("I. THÔNG TIN CHUNG")[2] < _match_heading("10. Điều kiện")[2]
    assert _match_heading("I. THÔNG TIN CHUNG")[3] == "roman-upper"


def test_bare_lowercase_i_stays_inside_the_letter_list():
    """"i)" inside an a) b) c)… list is the ninth item, not a new level."""
    assert _match_heading("i) Đơn phương hủy bỏ")[2] == _match_heading("h) Nội dung")[2]
    # while a bracketed "(i)" really is the deepest level
    assert _match_heading("(i) Đơn phương hủy bỏ")[2] > _match_heading("a) Nội dung")[2]


def test_letter_marker_keeps_its_separator():
    """An item marker keeps its punctuation ("c.", never bare "c")."""
    assert _match_heading("c. Ngày đáo hạn hợp đồng")[0] == "c."
    assert _match_heading("a) Hội viên gắn kết")[0] == "a)"
    assert _match_heading("(i) Đơn phương hủy bỏ")[0] == "(i)"


def test_heading_line_not_duplicated_in_chunk(processed):
    """A heading already in the prefix is not repeated on the line below."""
    for _path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            lines = [l for l in chunk.text.split("\n") if l.strip()]
            if len(lines) >= 2:
                assert lines[0].strip() != lines[1].strip(), chunk.chunk_id


def test_shorten_title_keeps_meaning():
    long_title = ("Sản phẩm áp dụng: Tiết kiệm thường, Tiết kiệm lẻ ngày, "
                  "Tiết kiệm bậc thang, Tiền gửi có kỳ hạn và nhiều loại khác nữa")
    assert shorten_title(long_title) == "Sản phẩm áp dụng"
    assert shorten_title("Hội viên gắn kết") == "Hội viên gắn kết"
    assert len(shorten_title("x" * 300)) <= 91


def test_flattened_table_row_names_the_section_by_its_content():
    """Flattened table row: the name is the meaningful value, not a column label.

    An external preprocessing step spreads each table row onto one line,
    "STT: 3.1 | Nội dung: Đối tượng khách hàng | Cụ thể: …". Cutting at the first
    colon, as for an ordinary heading, names every row "STT" and the outline can
    no longer tell one section from another.
    """
    row = "STT: 3.1 | Nội dung: Đối tượng khách hàng | Cụ thể:"
    assert shorten_title(row) == "Đối tượng khách hàng"
    # the index column is skipped wherever it appears
    assert shorten_title("Nội dung: Điều kiện vay | STT: 3.2") == "Điều kiện vay"
    # only when no short column is left does a long content column get used
    long_row = "STT: 3.5 | Nội dung: | Cụ thể: Lãi suất theo quy định của Ngân hàng"
    assert shorten_title(long_row) == "Lãi suất theo quy định của Ngân hàng"
    # a continuation row with every column blank -> no name, only the number
    assert shorten_title("STT: | Nội dung: | Cụ thể:") == ""
    # an ordinary heading containing a pipe still follows the original rule
    assert shorten_title("Hợp đồng bảo hiểm: là văn bản thỏa thuận") == "Hợp đồng bảo hiểm"


def test_tree_stops_at_level_three():
    """From level 4 ("1.1.1.1") down, headings leave the tree and become content."""
    lines = ["1. Phạm vi", "1.1 Đối tượng", "1.1.1 Cá nhân",
             "1.1.1.1 Đủ 18 tuổi trở lên", "1.1.1.1.1 Sâu hơn nữa", "1.2 Loại tiền"]
    blocks = [Block(text=t, page=1, bold=True) for t in lines]
    sections = [s for s in build_sections(detect_headings(blocks, "pdf"))
                if not s.is_preamble]
    assert [(s.level, s.heading_str) for s in sections] == [
        (1, "1 Phạm vi"), (2, "1.1 Đối tượng"),
        (3, "1.1.1 Cá nhân"), (2, "1.2 Loại tiền"),
    ]
    # the level 4/5 lines stay as content of the level 3 section, numbers intact
    body = [b.text for b in sections[2].blocks]
    assert body == ["1.1.1.1 Đủ 18 tuổi trở lên", "1.1.1.1.1 Sâu hơn nữa"]


def test_demoted_heading_keeps_its_number_from_word_numbering():
    """DOCX keeps the number outside the text: demote without splicing and it is lost."""
    from docindex.headings import _demote

    block = Block(text="Đủ 18 tuổi trở lên", page=1, number="1.1.1.1", kind="heading")
    _demote(block)
    assert block.text == "1.1.1.1 Đủ 18 tuổi trở lên"
    assert block.kind == "para" and block.number is None
    # in a PDF the number is already in the text and must not be spliced twice
    pdf_block = Block(text="1.1.1.1 Đủ 18 tuổi", page=1, number="1.1.1.1", kind="heading")
    _demote(pdf_block)
    assert pdf_block.text == "1.1.1.1 Đủ 18 tuổi"


def test_merging_a_short_section_adds_up_the_index():
    """Merging 1.2 into 1.1 adds up the numbers, it does not merge content alone."""
    from docindex.headings import merge_short_sections

    def sect(number, title, body):
        return Section(number=number, title=title, level=2,
                       path=["1. Gốc", f"{number} {title}"],
                       full_heading=f"{number} {title}",
                       blocks=[Block(text=body, page=1)], page_start=1, page_end=1)

    long_body = "nội dung dài " * 60
    out = merge_short_sections(
        [sect("1.1", "Mục đích", long_body), sect("1.2", "Loại tiền", "ngắn")],
        min_tokens=60, max_tokens=512)
    assert len(out) == 1
    assert out[0].number == "1.1 + 1.2"
    assert out[0].title == "Mục đích + Loại tiền"
    assert out[0].path == ["1. Gốc", "1.1 + 1.2 Mục đích + Loại tiền"]
    assert "ngắn" in " ".join(b.text for b in out[0].blocks), "absorbed content lost"


def test_short_section_stays_separate_when_the_merge_would_bust_the_limit():
    """A merge that would exceed 512 tokens must leave the sections apart."""
    from docindex.headings import merge_short_sections

    def sect(number, body):
        return Section(number=number, title="T", level=2,
                       path=["1. Gốc", f"{number} T"], full_heading=f"{number} T",
                       blocks=[Block(text=body, page=1)], page_start=1, page_end=1)

    almost_full = "chữ " * 470        # ~500 tokens; anything more overflows
    out = merge_short_sections([sect("1.1", almost_full), sect("1.2", "ngắn")],
                               min_tokens=60, max_tokens=512)
    assert [s.number for s in out] == ["1.1", "1.2"]


def test_merge_room_counts_the_path_prefix():
    """The 512 ceiling is the chunk's, and a chunk also carries the outline prefix.

    Forget the prefix and the merged section is split in two by the chunker
    anyway — the merge achieves nothing and produces two chunks with one name.
    """
    from docindex.headings import merge_short_sections

    long_path = ["PHẦN 1 " + "X" * 200, "ĐIỀU 1 " + "Y" * 200]

    def sect(number, body):
        return Section(number=number, title="T", level=3,
                       path=long_path + [f"{number} T"], full_heading=f"{number} T",
                       blocks=[Block(text=body, page=1)], page_start=1, page_end=1)

    # 330 + 60 = 390 < 512 ignoring the prefix, but the prefix eats ~150 tokens
    out = merge_short_sections([sect("1.1", "chữ " * 310), sect("1.2", "chữ " * 55)],
                               min_tokens=200, max_tokens=512)
    assert [s.number for s in out] == ["1.1", "1.2"], "merging would only be split again"


def test_merged_titles_do_not_bloat_the_path():
    """Merging many sections keeps every number; names are capped so the prefix stays small."""
    from docindex.headings import MAX_TITLE_IN_PATH, merge_short_sections

    sections = [
        Section(number=f"1.{i}", title=f"Tên mục khá dài số {i} của tài liệu",
                level=2, path=["1. Gốc", f"1.{i} T"], full_heading=f"1.{i} T",
                blocks=[Block(text="ngắn", page=1)], page_start=1, page_end=1)
        for i in range(1, 7)
    ]
    out = merge_short_sections(sections, min_tokens=60, max_tokens=512)
    assert len(out) == 1
    assert out[0].number == "1.1 + 1.2 + 1.3 + 1.4 + 1.5 + 1.6", "every number must be kept"
    assert len(out[0].title) <= MAX_TITLE_IN_PATH + 4
    assert out[0].title.endswith("…")


def test_flattened_row_with_a_missing_column_label():
    """When flattening loses a column's label, the whole cell is the value."""
    assert shorten_title("Khoản: 3.1 | Điều kiện vay vốn | Quy định:") == "Điều kiện vay vốn"
    assert (shorten_title("Khoản: 3.13 | Điều kiện tái cấp Hạn mức | Quy định: a. Trước")
            == "Điều kiện tái cấp Hạn mức")


def test_continuation_row_is_not_named_after_its_column():
    """A continuation row down to one column: "Quy định" is a label, not a name."""
    rows = [
        "3.1.1 Khoản: 3.1 | Điều kiện: Mục đích cho vay | Quy định: Cho vay bổ sung",
        "3.1.2 Quy định:",
        "3.1.3 Quy định: năm/lần ĐVKD thực hiện đánh giá lại hạn mức để xem xét lại mức",
        "3.1.4 Quy định: Việt Nam đồng",
    ]
    blocks = [Block(text=t, page=1, bold=True) for t in rows]
    sections = [s for s in build_sections(detect_headings(blocks, "pdf"))
                if not s.is_preamble]
    assert [s.title for s in sections] == [
        "Mục đích cho vay",     # a full row
        "",                     # an empty continuation row
        "",                     # a continuation row of overflow text, with no name
        "Việt Nam đồng",        # a continuation row whose value is short enough to name it
    ]
    # with no flattened table around, "Quy định:" is an ordinary section name
    plain = [Block(text="3.1.2 Quy định:", page=1, bold=True)]
    plain_sections = build_sections(detect_headings(plain, "pdf"))
    assert plain_sections[0].title == "Quy định"


def test_appendix_code_in_a_reference_is_not_a_heading():
    """"…theo Phụ lục số PL01.1016.PDS.2026(1);" is a reference, not a section."""
    assert _match_heading("PL01.1016.PDS.2026(1);") is None
    hit = _match_heading("PL02.1003.PCS.2026(1): Giải thích từ ngữ")
    assert hit[0] == "PL02.1003.PCS.2026(1)" and hit[1] == "Giải thích từ ngữ"


def test_numbered_appendix_is_not_named_after_itself():
    """"Phụ lục 01" has no name of its own — the path must not double it up."""
    blocks = [Block(text="Phụ lục 01", page=1, bold=True)]
    section = build_sections(detect_headings(blocks, "docx"))[0]
    assert section.title == ""
    assert section.heading_str == "Phụ lục 01"
    assert section.path == ["Phụ lục 01"]


def test_flattened_row_heading_line_carries_only_the_name():
    """On rebuild the heading line carries only the name; the row moves to the body."""
    from docindex.layout import _heading_and_rest

    row = "STT: 3.1 | Nội dung: Đối tượng khách hàng | Cụ thể:"
    section = Section(number="3.1.1", title=shorten_title(row), level=3,
                      path=["3.1.1"], full_heading=f"3.1.1 {row}")
    head, rest = _heading_and_rest(section)
    assert head == "3.1.1 Đối tượng khách hàng"
    # the original row keeps all its text in the body — every column keeps its label
    assert rest.startswith("STT: 3.1 | Nội dung: Đối tượng khách hàng")


def test_full_heading_is_not_truncated_in_content(processed):
    """The shortened heading is for the path only; the body keeps every word."""
    for _path, _chunks, sections, _src, _st in processed:
        for section in sections:
            # For a section with merged numbers, heading_str collects several
            # names while full_heading stays its own original line — the absorbed
            # section's words live in the body, which test_no_content_loss checks.
            if section.full_heading and not section.is_merged:
                assert len(section.full_heading) >= len(section.heading_str) - 1


def test_rows_to_markdown_drops_empty_columns():
    rows = [["", "Năm hợp đồng", "", "1", ""], ["", "Lãi suất", "", "3%", ""]]
    md = rows_to_markdown(rows)
    assert md.split("\n")[0] == "| Năm hợp đồng | 1 |"
    assert rows_to_markdown([["", ""], ["", ""]]) == ""


def test_long_table_row_keeps_all_cell_text():
    """A long row becomes 'column: value' lines without losing any text."""
    long_cell = "giá trị rất dài " * 90
    md = "\n".join([
        "| Tiêu chí | Nội dung |",
        "| --- | --- |",
        f"| Điều kiện | {long_cell} |",
    ])
    parts = _split_table(md, 800)
    joined = " ".join(parts)
    assert "Điều kiện" in joined
    assert joined.count("giá trị rất dài") == 90
    for part in parts:
        rows = [r for r in part.split("\n") if r.strip().startswith("|")]
        widths = {r.count("|") for r in rows}
        assert len(widths) <= 1, "the table's column count broke"


def test_toc_page_keeps_cover_content(processed):
    """Cover and table of contents on one page: only the TOC part may be dropped."""
    for path, chunks, _sections, _src, _st in processed:
        if "Tràng An" not in os.path.basename(path):
            continue
        first_page = " ".join(c.raw_text for c in chunks if c.page == 1)
        assert "AN LỘC TÍCH LŨY THỊNH VƯỢNG" in first_page
        assert "4474/BTC-QLBH" in first_page, "the approval document number was lost"
        assert "....." not in first_page, "a table-of-contents line survived"


def test_footnotes_are_dropped(processed):
    """A footnote is not content, and certainly not a node in the outline.

    "1 Theo địa giới hành chính cũ…", sitting under the rule at the foot of the
    page, annotates a point in the body. Kept, it looks exactly like heading
    "1." and is built into a major branch level with "1. Phạm vi điều chỉnh".
    """
    for path, chunks, sections, _src, _st in processed:
        if "Xuân Thành" not in os.path.basename(path):
            continue
        body = " ".join(c.raw_text for c in chunks)
        assert "Theo địa giới hành chính cũ" not in body
        assert "ĐVKD chủ động thẩm định và chịu trách nhiệm" not in body
        # The document's three headings, with no node born from a footnote line.
        # The first two are short and merge into one — both numbers are kept.
        titles = [s.title for s in sections]
        assert titles == [PREAMBLE_TITLE,
                          "Phạm vi điều chỉnh và đối tượng áp dụng + Giải thích từ ngữ",
                          "Nội dung sản phẩm"]
        assert [s.number for s in sections] == ["", "1 + 2", "3"]
        # a reference mark ("…từng thời kỳ⁶") must not stick to the word
        assert "thời kỳ6" not in body and "thời kỳ4" not in body
        assert "chuyển nợ về nhóm 1 thì" in body
        break
    else:
        pytest.fail("the Xuân Thành sample file is missing")


def _glyph_line(pieces, size=11.0, tracking=0.0):
    """Build a `rawdict`-style line: one cluster of glyphs per word, set flush.

    `tracking` is the width of the space glyph inserted between letters — a PDF
    flattening tool sets it to nearly zero, so it is invisible on paper, while a
    real space between words is 0.25 of the font size, as in any serif face.
    """
    chars = []
    x = 0.0

    def put(ch, width):
        nonlocal x
        chars.append({"c": ch, "bbox": (x, 0.0, x + width, size)})
        x += width

    for index, word in enumerate(pieces):
        if index:
            put(" ", size * 0.25)
        for pos, ch in enumerate(word):
            if pos and tracking:
                put(" ", tracking)
            put(ch, size * 0.5)
    return {"dir": (1.0, 0.0), "spans": [{"size": size, "chars": chars}]}


def test_ghost_spaces_between_letters_are_dropped():
    """A zero-width space between letters must not become a real space.

    A "flattened" PDF positions every glyph individually and inserts a space of
    almost no width between letters. Read straight out, the sentence shatters
    into "B á n , x á c" — chunks, tokenizer and rebuild all break with it.
    """
    from docindex.extract_pdf import _rebuild_line

    line = _glyph_line(["Bán,", "xác", "nhận"], tracking=0.2)
    _rebuild_line(line)
    assert line["spans"][0]["text"] == "Bán, xác nhận"


def test_real_spaces_survive_the_ghost_space_filter():
    """A real space stays a space — even when it closes out a span.

    Accented letters usually take their glyphs from a different font, so PyMuPDF
    cuts the line into several spans and a space often lands at the end of one
    ("QUY " | "ĐỊNH"). Measure the gap within a single span and the two words
    run together.
    """
    from docindex.extract_pdf import _rebuild_line

    line = _glyph_line(["QUY", "ĐỊNH", "NGÂN"])
    chars = line["spans"][0]["chars"]
    # cut right before "Đ": the space ahead of it closes the first span
    cut = next(i for i, c in enumerate(chars) if c["c"] == "Đ")
    line["spans"] = [{"size": 11.0, "chars": chars[:cut]},
                     {"size": 11.0, "chars": chars[cut:]}]
    _rebuild_line(line)
    assert "".join(s["text"] for s in line["spans"]) == "QUY ĐỊNH NGÂN"


def test_no_letter_spaced_runs_survive(processed):
    """No chunk still holds a letter-spaced run like "B á n , x á c n h ậ n"."""
    spaced = re.compile(r"(?:\w ){6,}\w")
    for path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            found = spaced.search(chunk.raw_text)
            assert not found, f"{os.path.basename(path)}: {found.group(0)!r}"


def test_flattened_prose_is_not_sliced_into_a_table(processed):
    """Prose must not be read as a table.

    The "Giải thích từ ngữ" page of the flattened file holds no table at all —
    only a numbered list of definitions. Guess the table frame from whitespace
    and it is sliced lengthwise into ten columns, with words broken across cells.
    """
    for path, chunks, sections, _src, _st in processed:
        if "bán ngoại tệ tiền mặt_flattened" not in os.path.basename(path):
            continue
        blocks = [b for s in sections for b in s.blocks]
        assert not [b for b in blocks
                    if b.is_table and "Mức phải khai báo" in b.text]
        body = re.sub(r"\s+", " ", " ".join(c.raw_text for c in chunks))
        assert ("Là mức ngoại tệ tiền mặt hoặc Việt Nam đồng tiền mặt của cá "
                "nhân mang theo khi xuất cảnh") in body
        break
    else:
        pytest.fail("the flattened sample file is missing")


def test_footnote_zone_leaves_real_content_alone(processed):
    """Only small text at the foot of the page is cut — body-size content stays."""
    for path, chunks, _sections, _src, _st in processed:
        if "Xuân Thành" not in os.path.basename(path):
            continue
        body = " ".join(c.raw_text for c in chunks)
        # the last body line on the page that carries a footnote
        assert "không quá 70 tuổi tại thời điểm kết thúc khoản vay" in body
        assert "Các giấy tờ khác có giá trị tương đương" in body
        break


def test_no_content_loss(processed):
    """The backstop against text loss: every word in the document reaches a chunk.

    Except table-of-contents lines, repeated footers and footnotes — three kinds
    of noise dropped on purpose. The 95% threshold leaves room for those.
    """
    import re
    from collections import Counter

    import docx
    import fitz

    from docindex import extract_pdf

    dot = re.compile(r"\.{4,}")

    def words(s):
        return re.findall(r"\w+", clean_text(s).lower())

    def pdf_lines(path):
        """PDF text, minus the footnote block and the repeated header lines.

        Footnotes and headers/footers are dropped on purpose, so counting them as
        text-that-must-survive would raise false alarms. They are located here
        from their own *definition* — an unbroken run of small text at the foot
        of the page, and lines near the top edge repeating on nearly every page —
        rather than by calling into the code under test, so the check still
        catches every other kind of text loss.

        The source is read through `_page_dict` rather than `get_text("dict")`:
        a flattened file carries ghost spaces between letters, so the raw text
        yields "b", "á", "n" as three separate "words". Use that as the baseline
        and every correctly reassembled sentence counts as lost text. Only the
        glyph reading is borrowed; the noise detection below is done here.
        """
        doc = fitz.open(path)
        sizes = Counter()
        pages = []
        for page in doc:
            rows = []
            for block in extract_pdf._page_dict(page)["blocks"]:
                if block["type"] != 0:
                    continue
                for line in block["lines"]:
                    spans = [s for s in line["spans"] if s["text"].strip()]
                    if not spans:
                        continue
                    text = "".join(s["text"] for s in spans)
                    size = round(max(s["size"] for s in spans), 1)
                    rows.append((line["bbox"][1] / page.rect.height, size, text))
                    sizes[size] += len(text)
            pages.append(sorted(rows))
        doc.close()

        body = sizes.most_common(1)[0][0] if sizes else 12.0

        def in_margin(y):
            return y <= 0.10 or y >= 0.90

        seen = Counter()
        for rows in pages:
            seen.update({text for y, _size, text in rows if in_margin(y)})
        repeated = {text for text, count in seen.items()
                    if count >= max(2, len(pages) // 2)}

        out = []
        for rows in pages:
            while rows and rows[-1][0] > 0.70 and rows[-1][1] <= body - 1.0:
                rows.pop()
            out += [text for y, _size, text in rows
                    if not (in_margin(y) and text in repeated)]
        return out

    for path, chunks, _sections, _src, _st in processed:
        lines = []
        if path.lower().endswith(".pdf"):
            lines = pdf_lines(path)
        else:
            document = docx.Document(path)
            lines = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    seen = set()
                    for cell in row.cells:
                        if cell._tc in seen:
                            continue
                        seen.add(cell._tc)
                        lines.append(cell.text)

        src = Counter()
        for line in lines:
            if dot.search(line):
                continue
            # Compare after symbol folding: Equation Editor's "𝐋𝟏" and "L1" are
            # the same text, differing in encoding rather than being lost.
            src.update(words(normalize_symbols(line)))
        got = Counter()
        for chunk in chunks:
            got.update(words(chunk.raw_text))

        total = sum(src.values())
        missing = sum((src - got).values())
        kept = (total - missing) / total if total else 1
        assert kept >= 0.95, (
            f"{os.path.basename(path)}: only {kept:.1%} of words kept, "
            f"{missing}/{total} missing"
        )


# --- writing the cleaned document ----------------------------------------

@pytest.fixture(scope="module")
def cleaned(tmp_path_factory):
    """Clean the whole real corpus once, shared by the tests below."""
    from docindex.export import clean_document, out_dir_for

    out = tmp_path_factory.mktemp("cleaned")
    results = []
    for path in _docs():
        # mirror the folder tree: two subfolders hold files with the same name
        dst_dir = out_dir_for(path, TEST_DIR, str(out))
        os.makedirs(dst_dir, exist_ok=True)
        dst, stats = clean_document(path, dst_dir)
        results.append((path, dst, stats))
    return results


def test_output_marks_the_file_name_as_formalized(cleaned):
    """The output name keeps the original name plus the `_formalized` suffix."""
    from docindex.export import FORMALIZED_SUFFIX

    for src, dst, _stats in cleaned:
        name, ext = os.path.splitext(os.path.basename(src))
        assert os.path.basename(dst) == f"{name}{FORMALIZED_SUFFIX}{ext.lower()}"


def test_output_does_not_repeat_the_suffix(tmp_path):
    """Re-running on an already-normalised file does not double the suffix."""
    from docindex.export import FORMALIZED_SUFFIX, out_path

    src = os.path.join(str(tmp_path), f"quy-dinh{FORMALIZED_SUFFIX}.pdf")
    dst = out_path(src, os.path.join(str(tmp_path), "ra"), out_format="same")
    assert os.path.basename(dst) == f"quy-dinh{FORMALIZED_SUFFIX}.pdf"


def test_output_refuses_to_overwrite_the_source(tmp_path):
    """The suffix never doubles, so a normalised file could overwrite itself."""
    import fitz

    from docindex.export import FORMALIZED_SUFFIX, clean_document, out_path

    src = os.path.join(str(tmp_path), f"quy-dinh{FORMALIZED_SUFFIX}.pdf")
    doc = fitz.open()
    doc.new_page()
    doc.save(src)
    doc.close()
    with pytest.raises(ValueError, match="is the source file"):
        clean_document(src, str(tmp_path))
    with pytest.raises(ValueError, match="is the source file"):
        out_path(src, str(tmp_path), out_format="same")


def test_cleaned_output_keeps_source_format(cleaned):
    for src, dst, _stats in cleaned:
        assert os.path.splitext(dst)[1].lower() == os.path.splitext(src)[1].lower()
        assert os.path.isfile(dst) and os.path.getsize(dst) > 0


def test_cleaned_pdf_keeps_content_figures(cleaned):
    """Figures stay in the file; they are not stripped along with the logos."""
    import fitz

    checked = 0
    for src, dst, stats in cleaned:
        if not src.lower().endswith(".pdf") or stats["figures_kept"] == 0:
            continue
        checked += 1
        doc = fitz.open(dst)
        drawn = 0
        for page in doc:
            for info in page.get_image_info():
                w = info["bbox"][2] - info["bbox"][0]
                h = info["bbox"][3] - info["bbox"][1]
                if min(w, h) >= 70:
                    drawn += 1
        doc.close()
        assert drawn >= stats["figures_kept"], (
            f"{os.path.basename(dst)}: {stats['figures_kept']} figures kept but "
            f"only {drawn} images large enough remain in the file"
        )
    assert checked, "no PDF with figures available to check"


def test_cleaned_pdf_drops_repeated_boilerplate(cleaned):
    """Repeated footers and page numbers must vanish from the cleaned file."""
    import re

    import fitz

    page_num = re.compile(r"^\s*\d+\s*/\s*\d+\s*$")
    for src, dst, _stats in cleaned:
        if not src.lower().endswith(".pdf"):
            continue
        doc = fitz.open(dst)
        for page in doc:
            for line in page.get_text("text").split("\n"):
                assert not page_num.match(line.strip()), (
                    f"{os.path.basename(dst)}: page number '{line.strip()}' survived"
                )
        doc.close()


def test_cleaned_pdf_is_not_bigger_than_source(cleaned):
    """Rewriting the PDF must not inflate it (it once nearly doubled)."""
    for src, dst, _stats in cleaned:
        if not src.lower().endswith(".pdf"):
            continue
        assert os.path.getsize(dst) <= os.path.getsize(src) * 1.05, (
            f"{os.path.basename(dst)} grew larger than the original"
        )


def test_cleaned_docx_keeps_body_and_clears_headers(cleaned):
    import docx

    for src, dst, _stats in cleaned:
        if not src.lower().endswith(".docx"):
            continue
        before = docx.Document(src)
        after = docx.Document(dst)
        assert sum(len(p.text) for p in after.paragraphs) == \
            sum(len(p.text) for p in before.paragraphs)
        assert len(after.tables) == len(before.tables)
        for section in after.sections:
            for part in (section.header, section.footer):
                assert not "".join(p.text for p in part.paragraphs).strip()


def test_faint_watermark_is_not_a_content_figure():
    """A watermark half a page wide is still classified as noise.

    It clears every size threshold and repeats on no page, so the only cue left
    is that it carries no dark ink at all.
    """
    import fitz

    from docindex.images import collect_pdf_images

    checked = 0
    for path in _docs():
        if "Trọn Đời" not in os.path.basename(path):
            continue
        checked += 1
        with fitz.open(path) as doc:
            found = [i for items in collect_pdf_images(doc).values() for i in items
                     if "watermark" in i.reason]
            assert found, "the watermark on page 3 was not recognised"
            assert all(i.kind == "logo" for i in found)
            # The other image on page 3 must be dropped for its own reason. This
            # check once demanded that page 3 keep a content figure, but that
            # image is a rounded red frame drawn under a paragraph — 452 live
            # characters of the page sit on top of it — so it is not a figure.
            for img in collect_pdf_images(doc).get(3, []):
                assert img.kind != "figure", f"{img.reason} is not a content figure"
    if not checked:
        pytest.skip("the Trọn Đời sample file is missing")


def test_decorative_frame_is_not_a_content_figure():
    """A frame drawn under the text is decoration; real diagrams must survive.

    The ink measure cannot separate the two: a frame is a few red strokes on a
    transparent ground, so nearly 100% of its non-paper pixels are dark ink and
    it scores 0.67, level with the densest diagram. What does separate them is
    the page's live text sitting on the frame, while a diagram's labels live
    inside the image file itself.
    """
    import fitz

    from docindex.images import collect_pdf_images

    frames = diagrams = 0
    for path in _docs():
        name = os.path.basename(path)
        if not path.lower().endswith(".pdf"):
            continue
        if "Trọn Đời" not in name and "Hưng Thịnh" not in name:
            continue
        with fitz.open(path) as doc:
            for items in collect_pdf_images(doc).values():
                for img in items:
                    if "decorative frame" in img.reason:
                        frames += 1
                        assert img.kind == "logo"
                    if img.kind == "figure":
                        diagrams += 1
    if not frames:
        pytest.skip("the Trọn Đời / Hưng Thịnh sample files are missing")
    # The "Tài khoản hợp đồng" diagram in Hưng Thịnh: its labels are inside the
    # image with no live text on top, so the decorative-frame rule must miss it.
    assert diagrams, "the content diagrams were stripped too"


def test_the_ink_threshold_sits_in_a_real_gap():
    """The ink threshold must sit inside a wide gap, close to neither side.

    One pale content figure landing near the threshold and the next tweak to the
    number would silently drop it. Measured on the real corpus: watermarks
    <= 0.02, content figures >= 0.17 — the 0.05 threshold is at least 3x from
    both.
    """
    import fitz

    from docindex.images import (
        WATERMARK_INK_RATIO, collect_pdf_images, dark_ink_ratio,
    )

    faint, solid = [], []
    for path in _docs():
        if not path.lower().endswith(".pdf"):
            continue
        with fitz.open(path) as doc:
            for pno, items in collect_pdf_images(doc).items():
                page = doc[pno - 1]
                for img in items:
                    if img.kind == "figure":
                        solid.append(dark_ink_ratio(page, img.bbox))
                    elif "watermark" in img.reason:
                        faint.append(dark_ink_ratio(page, img.bbox))
    assert solid, "no content figures available for comparison"
    assert min(solid) > WATERMARK_INK_RATIO * 3, (
        f"the palest content figure has only {min(solid):.3f} ink, near the threshold"
    )
    if faint:
        assert max(faint) < WATERMARK_INK_RATIO / 2, (
            f"the darkest watermark reaches {max(faint):.3f} ink, near the threshold"
        )


# --- choosing what to strip, item by item ---------------------------------

ONLY_IMAGES = dict(drop_logo=True, drop_cover=True,
                   drop_header_footer=False, drop_toc=False)


@pytest.fixture(scope="module")
def cleaned_images_only(tmp_path_factory):
    """The GUI's default level: strip only logos and cover art, leave text alone."""
    from docindex.export import CleanOptions, clean_document, out_dir_for

    out = tmp_path_factory.mktemp("images-only")
    results = []
    for path in _docs():
        dst_dir = out_dir_for(path, TEST_DIR, str(out))
        os.makedirs(dst_dir, exist_ok=True)
        dst, stats = clean_document(path, dst_dir, opts=CleanOptions(**ONLY_IMAGES))
        results.append((path, dst, stats))
    return results


def test_keeping_text_leaves_every_character_in_place(cleaned_images_only):
    """With text removal off, the cleaned file keeps every character of the source.

    That is what a user expects when only the image boxes are ticked: a lighter
    file that still opens page for page, line for line, like the original.
    """
    import fitz

    checked = 0
    for src, dst, stats in cleaned_images_only:
        if not src.lower().endswith(".pdf"):
            continue
        checked += 1
        assert stats["text_zones_removed"] == 0
        assert stats["pages_removed"] == 0
        with fitz.open(src) as before, fitz.open(dst) as after:
            assert after.page_count == before.page_count
            for old, new in zip(before, after):
                assert new.get_text("text") == old.get_text("text"), (
                    f"{os.path.basename(dst)} page {new.number + 1}: the text changed"
                )
    assert checked, "no PDF available to check"


def test_keeping_text_still_drops_the_logos(cleaned_images_only, cleaned):
    """Turning text removal off must not disable image removal as well."""
    full = {src: stats for src, _dst, stats in cleaned}
    dropped = 0
    for src, _dst, stats in cleaned_images_only:
        assert stats["images_removed"] == full[src]["images_removed"]
        dropped += stats["images_removed"]
    assert dropped, "no document had a logo to strip"


def test_keeping_images_leaves_them_in_the_file(tmp_path):
    """With both image boxes unticked, no image is stripped."""
    from docindex.export import CleanOptions, clean_document, out_dir_for

    opts = CleanOptions(drop_logo=False, drop_cover=False,
                        drop_header_footer=False, drop_toc=False)
    checked = 0
    for path in _docs():
        dst_dir = out_dir_for(path, TEST_DIR, str(tmp_path))
        os.makedirs(dst_dir, exist_ok=True)
        _dst, stats = clean_document(path, dst_dir, opts=opts)
        assert stats["images_removed"] == 0
        checked += 1
    assert checked


def test_cleaned_docx_can_keep_its_header_text(tmp_path):
    """Strip the logo from the header while the header text survives.

    The company logo is almost always in the header, so "logos only" has to
    touch the header without taking the text there along with it.
    """
    import docx

    from docindex.export import CleanOptions, clean_docx

    src = os.path.join(str(tmp_path), "co-header.docx")
    document = docx.Document()
    document.add_paragraph("Thân bài giữ nguyên.")
    document.sections[0].header.paragraphs[0].text = "Ban hành kèm Quyết định 01"
    document.save(src)

    dst = os.path.join(str(tmp_path), "ra.docx")
    clean_docx(src, dst, CleanOptions(**ONLY_IMAGES))
    after = docx.Document(dst)
    header = "".join(p.text for p in after.sections[0].header.paragraphs)
    assert "Ban hành kèm Quyết định 01" in header
    assert "Thân bài giữ nguyên." in "".join(p.text for p in after.paragraphs)


# --- rebuilt layout: one page per section ---------------------------------

@pytest.fixture(scope="module")
def laid_out(processed):
    from docindex.layout import build_pages

    return [(path, build_pages(sections), sections)
            for path, _chunks, sections, _src, _st in processed]


@pytest.fixture(scope="module")
def laid_out_per_section(processed):
    """The older layout: one page per section, each page within the token ceiling."""
    from docindex.layout import build_pages

    return [(path, build_pages(sections, page_per_section=True), sections)
            for path, _chunks, sections, _src, _st in processed]


def test_page_belongs_to_exactly_one_section(laid_out_per_section):
    """The layout's core constraint: no page holds content from two sections.

    A page may carry its own section's heading, or an ancestor's — an ancestor
    holds nothing but its heading line, so there is no content to mix in.
    """
    from docindex.layout import _heading_and_rest

    for path, pages, sections in laid_out_per_section:
        assert pages, os.path.basename(path)
        by_path = {s.path_str: s for s in sections}

        for page in pages:
            assert all(i.text.strip() for i in page.items)
            allowed = set()
            for depth in range(1, len(page.section.path) + 1):
                owner = by_path.get(" > ".join(page.section.path[:depth]))
                if owner is not None:
                    allowed.add(_heading_and_rest(owner)[0])

            for item in page.items:
                if item.kind != "heading":
                    continue
                name = item.text.removesuffix(" (tiếp)")
                assert name in allowed, (
                    f"{os.path.basename(path)}: heading '{name}' strayed onto the "
                    f"page of section {page.section.path_str}")


def test_section_content_stays_together(laid_out):
    """The parts of one section stay adjacent, with no other section between them."""
    for path, pages, _sections in laid_out:
        seen: list[str] = []
        for page in pages:
            key = page.section.path_str
            if not seen or seen[-1] != key:
                assert key not in seen, (
                    f"{os.path.basename(path)}: section {key} was split apart")
                seen.append(key)


def test_heading_size_decreases_by_level():
    """1. largest, 1.1 smaller, 1.1.1 smallest — and body smaller than all of them."""
    from docindex.layout import BODY_PT, HEADING_PT, heading_pt

    sizes = [heading_pt(level) for level in range(1, len(HEADING_PT) + 2)]
    assert sizes[0] > sizes[1] > sizes[2]
    assert all(a >= b for a, b in zip(sizes, sizes[1:]))
    assert min(sizes) > BODY_PT


def test_heading_is_on_its_own_line(laid_out):
    for _path, pages, _sections in laid_out:
        for page in pages:
            for item in page.items:
                if item.kind == "heading":
                    assert "\n" not in item.text.strip()


def test_long_section_is_split_evenly(laid_out_per_section):
    """A long section splits into roughly equal parts, none of them overflowing.

    A page is bounded twice: by the paper's capacity (lines) and by the RAG
    token ceiling. Whichever binds first decides the split, so balance is only
    demanded on the binding measure — text density per line varies, so the other
    measure being uneven is normal. Tables and figures cannot be split further,
    so the lightest part can fall below an even share by at most the largest
    block; more than that means the algorithm is filling greedily rather than
    dividing evenly.
    """
    from docindex.layout import LINES_PER_PAGE

    def balanced(sizes: list[int], biggest: int) -> bool:
        return min(sizes) >= sum(sizes) / len(sizes) - biggest

    checked = 0
    for path, pages, _sections in laid_out_per_section:
        groups: dict[str, list] = {}
        for page in pages:
            if page.part_total > 1:
                groups.setdefault(page.section.path_str, []).append(page)
        for key, parts in groups.items():
            if len(parts) < 2:
                continue
            checked += 1
            lines = [p.lines for p in parts]
            tokens = [p.tokens for p in parts]
            assert max(lines) <= LINES_PER_PAGE, (
                f"{os.path.basename(path)} {key}: page overflows at {lines} lines")
            assert (balanced(lines, max(i.lines for p in parts for i in p.items))
                    or balanced(tokens, max(i.tokens for p in parts for i in p.items))), (
                f"{os.path.basename(path)} {key}: parts are uneven at "
                f"{lines} lines / {tokens} tokens")
    assert checked, "no section longer than one page available to check"


def test_page_fits_the_rag_token_limit(laid_out_per_section):
    """Every page fits inside 512 tokens.

    The RAG stack reads the document by its DLA tree and chunks at a 512 token
    ceiling. A page over that limit gets cut at an arbitrary point mid-sentence,
    so the split is made at a sentence boundary during the rebuild instead.
    """
    from docindex.chunker import MAX_TOKENS

    for path, pages, _sections in laid_out_per_section:
        for page in pages:
            biggest = max(i.tokens for i in page.items)
            if biggest > MAX_TOKENS:
                continue        # a single block that cannot be split any further
            assert page.tokens <= MAX_TOKENS, (
                f"{os.path.basename(path)} {page.section.path_str}: "
                f"page {page.part_index} weighs {page.tokens} tokens")


def test_rebuilt_pdf_shows_headings_as_titles(processed, tmp_path_factory):
    """Headings in the PDF are bold and larger than the body text.

    A DLA model looks at the printed page, not the file structure: size and
    weight are what decide whether it labels a line `title` or `list-item`.
    """
    import fitz

    from docindex.export import rebuild_document
    from docindex.layout import BODY_PT, build_pages

    out = tmp_path_factory.mktemp("dla")
    path, _chunks, sections, _src, _st = processed[0]
    heads = [i for p in build_pages(sections) for i in p.items if i.kind == "heading"]
    dst, _stats = rebuild_document(sections, path, str(out), out_format="pdf")

    doc = fitz.open(dst)
    try:
        spans = [s for page in doc
                 for block in page.get_text("dict")["blocks"]
                 if block.get("type") == 0
                 for line in block["lines"] for s in line["spans"]]
    finally:
        doc.close()

    checked = 0
    for head in heads:
        probe = head.text.strip()[:12]
        # A short section name ("Phụ lục 01") also appears in cross-references in
        # the body, so *every* matching span is considered and the largest taken —
        # taking the first match would catch a body line instead.
        hits = [s for s in spans if probe and probe in s["text"]]
        if not hits:
            continue                # the heading was wrapped mid-line
        hit = max(hits, key=lambda s: s["size"])
        checked += 1
        assert hit["size"] > BODY_PT * 1.1, (
            f"heading '{probe}' is only {hit['size']}pt, no larger than the body")
        bold = bool(hit["flags"] & 2 ** 4) or "bold" in hit["font"].lower()
        assert bold, f"heading '{probe}' is not bold ({hit['font']})"
    assert checked >= len(heads) // 2, "no headings could be checked in the PDF"


def test_rebuilt_document_invents_no_heading(laid_out):
    """Every heading line drawn must actually exist in the document.

    A RAG stack builds its outline from the heading lines it sees, so a line the
    tool invented (the preamble label, say) is a phantom branch in the tree and
    a wrong title on every chunk below it.

    The one exception is the "(tiếp)" suffix: a section over the token ceiling
    has to be split into parts, and the line reopening it is exactly where the
    RAG stack makes the cut. It invents no new section — the name is intact in
    front of it.
    """
    from docindex.layout import _heading_and_rest
    from docindex.models import document_title

    for path, pages, sections in laid_out:
        real = {_heading_and_rest(s)[0] for s in sections if not s.is_preamble}
        real.add(document_title(path))
        for page in pages:
            for item in page.items:
                if item.kind != "heading":
                    continue
                name = item.text[:-len(" (tiếp)")] if item.text.endswith(
                    " (tiếp)") else item.text
                assert name in real, (
                    f"{os.path.basename(path)}: heading '{item.text}' is not in "
                    "the original document")


def test_every_level_keeps_one_font_size(laid_out):
    """One level, one font size, and no deeper level is larger than a shallower one.

    A DLA model infers heading level from font size. If one level is sometimes
    large and sometimes small, the outline it builds comes out ragged.
    """
    sizes: dict[int, set[float]] = {}
    for _path, pages, _sections in laid_out:
        for page in pages:
            for item in page.items:
                if item.kind == "heading":
                    sizes.setdefault(item.level, set()).add(item.size)

    assert sizes
    for level, values in sizes.items():
        assert len(values) == 1, f"level {level} uses several font sizes: {values}"
    ordered = [next(iter(sizes[level])) for level in sorted(sizes)]
    assert ordered == sorted(ordered, reverse=True), ordered


def test_headings_ask_word_to_keep_them_with_their_content(processed, tmp_path_factory):
    """Headings in the .docx carry the keep_with_next flag.

    Word decides the page breaks and the tool does not interfere. This flag is
    what keeps a heading from being stranded at the foot of a page — without it,
    a heading and its content end up on two different pages.
    """
    import docx
    from docx.shared import Pt

    from docindex.export import rebuild_document
    from docindex.layout import BODY_PT

    out = tmp_path_factory.mktemp("keep")
    path, _chunks, sections, _src, _st = processed[0]
    dst, _stats = rebuild_document(sections, path, str(out), out_format="docx")

    heads = [p for p in docx.Document(dst).paragraphs
             if p.runs and p.runs[0].bold and p.runs[0].font.size > Pt(BODY_PT)]
    assert heads, "no headings found in the .docx"
    for para in heads:
        assert para.paragraph_format.keep_with_next, (
            f"heading '{para.text[:40]}' is missing keep_with_next")


def test_outline_stays_within_the_rag_depth_limit(processed):
    """The RAG stack reads an outline at most 6 levels deep."""
    from docindex.layout import MAX_TREE_DEPTH

    for path, _chunks, sections, _src, _st in processed:
        depth = max(s.level for s in sections)
        assert depth <= MAX_TREE_DEPTH, (
            f"{os.path.basename(path)}: outline is {depth} levels deep")


def test_rebuilt_pdf_leads_with_the_document_title(processed, tmp_path_factory):
    """Page one opens with the document title, bold and the largest on the page.

    A RAG stack uses the document title as the root of every chunk title;
    without it, every chunk loses its topmost level.
    """
    import fitz

    from docindex.export import rebuild_document
    from docindex.models import document_title

    out = tmp_path_factory.mktemp("title")
    path, _chunks, sections, _src, _st = processed[0]
    dst, stats = rebuild_document(sections, path, str(out), out_format="pdf")
    assert stats["doc_title"] == document_title(path)

    doc = fitz.open(dst)
    spans = [s for block in doc[0].get_text("dict")["blocks"]
             if block.get("type") == 0
             for line in block["lines"] for s in line["spans"]]
    doc.close()

    biggest = max(spans, key=lambda s: s["size"])
    assert biggest["text"].strip()[:20] in stats["doc_title"]
    assert bool(biggest["flags"] & 2 ** 4) or "bold" in biggest["font"].lower()


def test_rebuilt_pdf_uses_plain_spaces(processed, tmp_path_factory):
    """Reading the PDF back yields ordinary spaces, not U+00A0.

    A TrueType font embedded in a PDF makes the space glyph map to U+00A0 in the
    ToUnicode table. Nothing looks different on paper, but every extraction tool
    receives a string without a single ordinary space in it — word segmentation,
    BM25 and the RAG tokenizer all break, and none of them report an error.
    """
    import fitz

    from docindex.export import rebuild_document

    out = tmp_path_factory.mktemp("spaces")
    path, _chunks, sections, _src, _st = processed[0]
    dst, _stats = rebuild_document(sections, path, str(out), out_format="pdf")

    doc = fitz.open(dst)
    text = "".join(page.get_text("text") for page in doc)
    doc.close()
    assert " " not in text, "spaces in the PDF are still U+00A0"
    assert text.count(" ") > 100, "no ordinary spaces were extracted"


def test_rebuilt_document_keeps_all_words(laid_out):
    """Rebuilding the layout drops no word of the extracted content."""
    import re
    from collections import Counter

    def words(s):
        return re.findall(r"\w+", clean_text(s).lower())

    for path, pages, sections in laid_out:
        src = Counter()
        for section in sections:
            if not section.is_preamble:
                # "Phần mở đầu" là tên do tool đặt, không có trong tài liệu
                src.update(words(section.heading_text))
            for block in section.blocks:
                if block.kind == "figure":
                    continue        # figures become real images, not text
                src.update(words(block.text))
        got = Counter()
        for page in pages:
            for item in page.items:
                got.update(words(item.text))

        missing = sum((src - got).values())
        total = sum(src.values())
        assert missing == 0, (
            f"{os.path.basename(path)}: {missing}/{total} words lost in the rebuild")


def test_rebuild_writes_both_formats(processed, tmp_path_factory):
    """Either .docx or .pdf; the reported page count matches the real file."""
    import docx
    import fitz
    from docx.oxml.ns import qn

    from docindex.export import rebuild_document
    from docindex.layout import build_pages

    out = tmp_path_factory.mktemp("rebuilt")
    path, _chunks, sections, _src, _st = processed[0]

    # Continuous-flow mode: the renderer decides the page breaks, so the document
    # carries none of its own; forcing breaks from the tool's estimate is what
    # produced the nearly-empty pages.
    dst_docx, stats = rebuild_document(sections, path, str(out), out_format="docx")
    assert dst_docx.endswith(".docx") and stats["pages_out"] >= 1
    document = docx.Document(dst_docx)
    breaks = sum(1 for p in document.paragraphs for br in p._p.iter(qn("w:br"))
                 if br.get(qn("w:type")) == "page")
    assert breaks == 0, "the continuous-flow build must insert no page breaks"

    dst_pdf, stats = rebuild_document(sections, path, str(out), out_format="pdf")
    doc = fitz.open(dst_pdf)
    assert dst_pdf.endswith(".pdf") and stats["pages_out"] == doc.page_count

    # In one-page-per-section mode the break count matches the pages laid out
    per_section = len(build_pages(sections, page_per_section=True))
    dst2, stats2 = rebuild_document(sections, path, str(out), out_format="docx",
                                    suffix="_persection", page_per_section=True)
    breaks2 = sum(1 for p in docx.Document(dst2).paragraphs
                  for br in p._p.iter(qn("w:br")) if br.get(qn("w:type")) == "page")
    assert breaks2 == per_section - 1 and stats2["pages_out"] == per_section
    # the font must carry the Vietnamese diacritics, not fall back to tofu boxes
    text = "".join(page.get_text("text") for page in doc)
    doc.close()
    assert any(ch in text for ch in "ăâđêôơư"), "the PDF lost its Vietnamese diacritics"


def _docx_with_picture(tmp_path):
    """Build a .docx with an embedded image — the real corpus has no such file."""
    import docx
    import fitz
    from docx.shared import Pt as DocxPt

    png = tmp_path / "hinh.png"
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 260, 180))
    pix.clear_with(210)
    pix.save(str(png))

    document = docx.Document()
    document.add_paragraph("1. Mục có hình minh hoạ")
    document.add_picture(str(png), width=DocxPt(260))
    src = tmp_path / "co-hinh.docx"
    document.save(str(src))
    return str(src)


def test_docx_image_is_extracted_and_rebuilt(tmp_path):
    """A figure in a .docx exports to a file so the rebuild can embed it again."""
    import docx
    import fitz

    from docindex.export import rebuild_document
    from docindex.headings import build_sections, detect_headings
    from docindex.extract_docx import extract

    src = _docx_with_picture(tmp_path)
    figure_dir = str(tmp_path / "figures")
    blocks, _source = extract(src, figure_dir=figure_dir, doc_id="t")
    figures = [b for b in blocks if b.kind == "figure"]
    assert figures, "no figure was recognised in the .docx"
    assert os.path.isfile(figures[0].meta["file"]), "the image file was not extracted"

    sections = build_sections(detect_headings(blocks, "docx"))
    out = str(tmp_path / "out")
    dst, _stats = rebuild_document(sections, src, out, out_format="docx",
                                   figure_dir=figure_dir)
    parts = docx.Document(dst).part.package.image_parts
    assert len(list(parts)) == 1, "the image was not embedded in the rebuilt .docx"

    dst, _stats = rebuild_document(sections, src, out, out_format="pdf",
                                   figure_dir=figure_dir)
    doc = fitz.open(dst)
    embedded = sum(len(page.get_images()) for page in doc)
    doc.close()
    assert embedded == 1, "the image was not embedded in the rebuilt .pdf"


def test_rebuilt_pdf_is_not_bloated_by_fonts(processed, tmp_path_factory):
    """Embedding a whole system font once inflated the file by several MB."""
    from docindex.export import rebuild_document

    out = tmp_path_factory.mktemp("rebuilt-size")
    path, _chunks, sections, _src, _st = processed[0]
    dst, _stats = rebuild_document(sections, path, str(out), out_format="pdf")
    assert os.path.getsize(dst) < 1_500_000, "the font was not subset"


# --- graphical interface --------------------------------------------------

def test_outline_shows_the_tree_with_indentation(processed):
    """The printed outline indents by level and holds no tool-invented section."""
    from docindex.models import PREAMBLE_TITLE, document_title
    from docindex.report import format_outline, outline

    path, _chunks, sections, _src, _st = processed[0]
    tree = format_outline(sections, document_title(path))
    lines = tree.split("\n")

    assert lines[0].startswith("[document] ")
    assert PREAMBLE_TITLE not in tree
    nodes = outline(sections)
    assert len(lines) == len(nodes) + 1
    for line, node in zip(lines[1:], nodes):
        assert line.startswith("   " * node["level"] + f"L{node['level']} ")


def test_gui_options_cover_everything_the_worker_reads():
    """A missing key in opts kills the GUI halfway through a run."""
    import inspect

    from docindex import gui

    source = inspect.getsource(gui.App._work)
    used = set(re.findall(r"opts\[\"(\w+)\"\]", source))
    built = set(re.findall(r"\"(\w+)\":", inspect.getsource(gui.App.start)))
    assert used <= built, f"opts is missing keys: {used - built}"


def test_parse_drop_handles_paths_with_spaces():
    """The OS wraps paths containing spaces in braces."""
    from docindex.gui import parse_drop

    assert parse_drop("{C:/Máy tính/tài liệu a.pdf} C:/b.docx") == [
        "C:/Máy tính/tài liệu a.pdf", "C:/b.docx",
    ]
    assert parse_drop("C:/x.pdf") == ["C:/x.pdf"]
    assert parse_drop("") == []
    # Windows sends backslashes and braces each file when several are dropped
    assert parse_drop(r"{C:\Tài liệu\quy dinh.docx} {C:\a\b c.pdf}") == [
        r"C:\Tài liệu\quy dinh.docx", r"C:\a\b c.pdf",
    ]


def test_gui_says_so_when_drag_and_drop_is_unavailable():
    """Without tkinterdnd2, say so plainly rather than inviting a pointless drop.

    If the drop area still reads "Drop documents here" while the library is
    missing, the user drops a file, nothing happens, and there is no clue why.
    """
    import tkinter as tk

    from docindex import gui

    try:
        root = tk.Tk()          # plain Tk: drag and drop is definitely unavailable
    except tk.TclError:
        pytest.skip("this environment cannot open a Tk window")
    try:
        app = gui.App(root)
        root.update()
        assert not app._enable_dnd()
        assert "unavailable" in app.drop_label["text"]
        assert "Choose files" in app.drop_label["text"]
        assert "tkinterdnd2" in app.log.get("1.0", "end")
    finally:
        root.destroy()


def test_expand_inputs_filters_and_walks_folders():
    from docindex.gui import expand_inputs

    from_folder = expand_inputs([TEST_DIR])
    assert from_folder and all(
        os.path.splitext(p)[1].lower() in {".pdf", ".docx"} for p in from_folder
    )
    assert expand_inputs(["does-not-exist.txt"]) == []


def test_page_numbers_are_not_in_text(processed):
    """A footer line like '9/34' is noise and must be stripped entirely."""
    import re
    pat = re.compile(r"^\s*\d+\s*/\s*\d+\s*$")
    for _path, chunks, _sections, _src, _st in processed:
        for chunk in chunks:
            for line in chunk.raw_text.split("\n"):
                assert not pat.match(line), f"{chunk.chunk_id}: '{line}' survived"


# --- normalising special symbols ------------------------------------------

def test_math_variables_become_plain_letters():
    """Equation Editor variables are lowered to plain Latin letters.

    Interest formulas are typed in Equation Editor, so "Mi" really lives in the
    Mathematical Alphanumeric Symbols block. A RAG stack has no font for that
    block and reads it as "$Mi$" or drops it — the correct result is "Mi".
    """
    assert normalize_symbols("𝐌𝐢") == "Mi"
    assert normalize_symbols("𝐋 = 𝐌 ∗𝐓 ∗𝐑") == "L = M *T *R"
    assert normalize_symbols("𝟏𝟐") == "12"
    assert "$" not in clean_text("𝑳: Là số tiền lãi")


def test_symbols_are_spelled_out_or_flattened():
    """Maths symbols and typographic punctuation fold to words or ASCII marks."""
    assert "Tổng" in normalize_symbols("∑")
    assert normalize_symbols("a ≤ b ≥ c ≠ d") == "a <= b >= c != d"
    assert normalize_symbols("5 × 3 ÷ 2 ± 1") == "5 x 3 / 2 +/- 1"
    assert normalize_symbols("“abc” – ‘x’") == "\"abc\" - 'x'"
    # Word's Symbol font pushes glyphs into the PUA, keeping their ASCII codes
    assert normalize_symbols("a \uf02b b") == "a + b"
    assert normalize_symbols("\uf0b7 gạch đầu dòng") == "- gạch đầu dòng"


def test_clean_text_keeps_ordinary_vietnamese_intact():
    """Normalisation leaves ordinary Vietnamese text and punctuation untouched."""
    src = "Điều 5.1: Lãi suất 6,5%/năm (áp dụng từ 01/2026) - xem mục 2.3."
    assert clean_text(src) == src


def test_rebuilt_pdf_reads_back_every_character(tmp_path):
    """Reading the PDF back yields exactly the characters that were written.

    The ToUnicode table MuPDF builds maps a few glyphs to their twins — hyphen
    to soft hyphen U+00AD, semicolon to the Greek question mark U+037E. Nothing
    looks different on paper, but the RAG side receives a string with not one
    hyphen or semicolon left in it.
    """
    import fitz

    from docindex.layout import LayoutPage, PageItem
    from docindex.render import write_pdf

    probe = "a-b; c-d; Điều 5.1 - mục 2,3; x/y"
    dst = str(tmp_path / "probe.pdf")
    write_pdf([LayoutPage(None, [PageItem("para", probe)])], dst)

    doc = fitz.open(dst)
    text = "".join(page.get_text("text") for page in doc)
    doc.close()
    for mark in ("-", ";", " "):
        assert mark in text, f"reading the PDF back lost the {mark!r} character"
    for wrong in ("\u00ad", "\u037e", "\u00a0"):
        assert wrong not in text, f"the PDF reads back the twin character U+{ord(wrong):04X}"


# --- unnumbered banner headings -------------------------------------------

def test_two_big_titles_become_the_top_level(processed):
    """When a document has several banner headings, they are the tree's top level.

    "TỔNG QUAN VĂN BẢN QUY ĐỊNH" followed by "QUY ĐỊNH SẢN PHẨM…" divides the
    document into two major parts. Miss them and the first part's "1. TIÊU ĐỀ
    SẢN PHẨM" sits level with the second part's "Điều 1", and the outline loses
    its top level entirely.
    """
    for path, _chunks, sections, _src, _st in processed:
        if "rút gốc linh hoạt" not in os.path.basename(path):
            continue
        banners = [s for s in sections if s.is_banner]
        assert len(banners) >= 2, "no banner heading was recognised"
        assert all(s.level == 1 for s in banners), "banner headings are not at level 1"
        assert all(not s.number for s in banners), "a banner heading carries a number"
        assert "TỔNG QUAN VĂN BẢN QUY ĐỊNH" in [s.title for s in banners]
        # every numbered section sits under one banner heading or another
        numbered = [s for s in sections if s.number]
        assert numbered and all(s.level >= 2 for s in numbered)
        break
    else:
        pytest.fail("the 'rút gốc linh hoạt' sample file is missing")


def test_a_formula_line_is_not_taken_for_a_big_title():
    """"L = M *T *R" is also large and all caps, but it is not a heading."""
    from docindex.headings import _is_banner

    def block(text):
        return Block(text=text, page=1, size=16.0, bold=True,
                     meta={"body_size": 12.0})

    assert not _is_banner(block("L = M *T *R"), 12.0)
    assert not _is_banner(block("12"), 12.0)
    assert _is_banner(block("TỔNG QUAN VĂN BẢN QUY ĐỊNH"), 12.0)
    # at body font size it is not a banner heading
    assert not _is_banner(block("TỔNG QUAN VĂN BẢN"), 16.0)


# --- merging short sections -----------------------------------------------

def test_merging_looks_both_ways(processed):
    """A short section considers both the section above it and the one below.

    Looking only backwards leaves a short section stranded forever behind a huge
    one — a hundred-token "Điều 1" standing alone only because an
    eight-hundred-token "MỤC LỤC" precedes it, while "Điều 2" just below has
    room to spare.
    """
    from docindex.headings import merge_short_sections

    for path, _chunks, sections, _src, _st in processed:
        if "rút gốc linh hoạt" not in os.path.basename(path):
            continue
        assert [s for s in sections if s.is_merged], "no short section was merged"
        # a second pass must merge nothing more: the fixed point has been reached
        again = merge_short_sections(sections, CFG.min_tokens, CFG.max_tokens)
        assert len(again) == len(sections), "a mergeable pair of short sections is left"
        break
    else:
        pytest.fail("the 'rút gốc linh hoạt' sample file is missing")


def test_merged_section_does_not_repeat_its_own_name(laid_out):
    """A section with merged numbers must not reprint its own name below the heading.

    A glossary has *the section name as its content* ("12 Doanh nghiệp cho thuê
    lại lao động: Là doanh nghiệp…"). Merge two such sections and the name sits
    both on the heading line and at the start of the body — read back, it
    repeats itself word for word.
    """
    from docindex.layout import _heading_and_rest

    for _path, _pages, sections in laid_out:
        for section in sections:
            if not section.is_merged or not section.own_title:
                continue
            _title, rest = _heading_and_rest(section)
            first = rest.split("\n")[0].strip()
            assert not first.startswith(section.own_title), (
                f"'{section.own_title}' is both the heading and the first body line")


# --- full-length headings -------------------------------------------------

def test_heading_line_keeps_the_whole_name(laid_out):
    """The printed heading takes the name at full length, not the shortened form.

    Names are trimmed to 90 characters so the outline prefix leaves room for
    content. Use that trimmed form as the printed heading and "…làm nguyên" ends
    up on the heading while "liệu sản xuất" drops down as a stunted paragraph.
    """
    from docindex.headings import MAX_TITLE_IN_HEADING, MAX_TITLE_IN_PATH

    longer = 0
    for _path, _pages, sections in laid_out:
        for section in sections:
            if section.is_preamble or not section.own_title:
                continue
            # trimming is only allowed when the name truly exceeds the heading ceiling
            assert (not section.own_title.endswith("…")
                    or len(section.own_title) >= MAX_TITLE_IN_HEADING - 20), (
                f"section name '{section.own_title}' was trimmed too early")
            longer += len(section.own_title) > MAX_TITLE_IN_PATH
    assert longer, "no heading exceeds the path ceiling — the sample is too small"


def test_long_heading_keeps_every_word(processed):
    """A long heading in a legal text must not be dismissed as prose."""
    for path, _chunks, sections, _src, _st in processed:
        if "Ký quỹ" not in os.path.basename(path):
            continue
        target = [s for s in sections if s.number.startswith("15.10")]
        assert target, "heading 15.10 was lost entirely"
        assert "nguyên liệu sản xuất" in target[0].full_heading, (
            f"heading 15.10 is missing words: {target[0].full_heading!r}")
        break
    else:
        pytest.fail("the 'Ký quỹ' sample file is missing")


# --- over-long sections ---------------------------------------------------

def test_long_section_is_reopened_with_a_continuation_line(laid_out):
    """A section over the token ceiling is split with a "(tiếp)" line.

    The RAG stack chunks by the outline tree rather than by page, so a
    four-thousand-token section yields exactly one four-thousand-token chunk and
    is truncated at embedding time. The "(tiếp)" line is the node it splits on.
    """
    from docindex.chunker import MAX_TOKENS

    seen = 0
    for path, pages, _sections in laid_out:
        for page in pages:
            run = 0
            for item in page.items:
                if item.kind == "heading":
                    seen += item.text.endswith("(tiếp)")
                    run = 0
                    continue
                run += item.tokens
                assert run <= MAX_TOKENS * 2, (
                    f"{os.path.basename(path)}: a {run} token block has no "
                    "'(tiếp)' line splitting it")
    assert seen, "no document has a section long enough to split — sample too small"


# --- page breaks ----------------------------------------------------------

def test_no_heading_is_left_alone_at_the_bottom_of_a_page(processed, tmp_path_factory):
    """A heading is never stranded at the foot of a page, nor split across two.

    A DLA model reads the outline page by page. A heading at the foot of one
    page with its content on the next is read as an empty section, and the next
    page's content as a section with no name.
    """
    import fitz

    from docindex.export import rebuild_document
    from docindex.render import _stranded_pages

    out = tmp_path_factory.mktemp("breaks")
    checked = 0
    for path, _chunks, sections, _src, _st in processed[:6]:
        dst, _stats = rebuild_document(sections, path, str(out), out_format="pdf")
        doc = fitz.open(dst)
        try:
            stranded = _stranded_pages(doc)
        finally:
            doc.close()
        assert not stranded, (
            f"{os.path.basename(path)}: heading stranded on page "
            f"{sorted(p + 1 for p in stranded)}")
        checked += 1
    assert checked, "no document could be rebuilt for checking"


# --- vector diagrams ------------------------------------------------------

def test_flowchart_is_kept_as_a_picture(tmp_path):
    """A process flowchart stays a picture; it is not pulled apart into text.

    A flowchart is vector strokes plus the text inside each box. Ordinary
    extraction lays all that text into one meaningless line and the diagram
    itself disappears — losing the picture and adding noise at once.
    """
    import fitz

    from docindex.images import collect_vector_figures

    name = "Quy định phan phoi trai phieu LPBank ra công chúng.pdf"
    path = os.path.join(TEST_DIR, "Tiết kiệm kênh bank", name)
    if not os.path.isfile(path):
        pytest.skip("the sample file containing a flowchart is missing")

    doc = fitz.open(path)
    try:
        found = [page.number + 1 for page in doc if collect_vector_figures(page)]
    finally:
        doc.close()
    assert found, "no flowchart was recognised"

    chunks, _sections, _src = process_file(
        path, CFG, figure_dir=str(tmp_path / "figures"))
    body = " ".join(c.raw_text for c in chunks)
    assert "[FIGURE:" in body, "the flowchart left no placeholder line"
    # the text inside the flowchart's boxes must not be pulled out as content
    assert "Tiếp nhận và khai báo" not in body


def test_a_bordered_table_is_not_taken_for_a_flowchart():
    """Tables are drawn with strokes too, but all orthogonal — not diagrams."""
    import fitz

    from docindex.images import collect_vector_figures

    for path in _docs():
        if not path.lower().endswith(".pdf"):
            continue
        doc = fitz.open(path)
        try:
            for page in doc:
                for figure in collect_vector_figures(page):
                    # every figure found must be found for its diagonal or
                    # curved strokes
                    assert "diagonal/curved" in figure.reason
        finally:
            doc.close()
